"""
analytical_report.py - Generation du rapport analytique en 2 etapes.

ETAPE 1 (lourde) : build_report_content()
  - Passe IA 1 : plan d'analyse
  - Calculs pandas + graphes matplotlib (PNG en memoire)
  - Passe IA 2 : interpretations
  -> Retourne un dict "report_content" complet, serialisable
     (avec graphes encodes en base64 pour le frontend)

ETAPE 2 (legere) : compose_word_from_content(report_content)
  - Prend le contenu (potentiellement edite par l'utilisateur)
  - Compose le .docx
  -> Retourne les bytes

Cela permet : preview HTML cote frontend -> edition -> telechargement Word.

------------------------------------------------------------------------
v2 (refonte graphique) :
  - Histogramme + KDE (au lieu d'histogramme classique)
  - Donut chart avec total au centre (au lieu de pie)
  - Lollipop chart (au lieu de barres horizontales)
  - Heatmap (au lieu de stacked bar) pour cat x cat
  - Boxplot horizontal (au lieu de barres de moyennes) pour cat x num
  - Scatter + droite de regression + bande de confiance 95% (au lieu de scatter brut)
  - Palette etendue : navy + gold SISTA + 6 couleurs editoriales harmonieuses
  - Colormap sequentielle SISTA pour heatmaps
------------------------------------------------------------------------
"""

from __future__ import annotations

import io
import os
import json
import base64
from datetime import datetime
from typing import Any

import pandas as pd
import numpy as np

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
from matplotlib.patches import Rectangle

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core import ai_agent

# ----------------------------------------------------------------------
#  Charte graphique SISTA - palette modernisee
# ----------------------------------------------------------------------

SISTA = {
    "navy":         "#13263D",  # Navy primaire SISTA
    "navy_deep":    "#0D1B2C",  # Navy profond
    "navy_soft":    "#4A6FA5",  # Navy doux (steel blue)
    "gold":         "#EFC71A",  # Gold primaire SISTA
    "gold_deep":    "#D4AC0D",  # Gold profond
    "beige":        "#F4F7FA",  # Fond clair
    "cream":        "#FAF6EC",  # Fond creme (chaud)
    "gray":         "#6B7280",  # Gris texte
    "gray_light":   "#E5E7EB",  # Gris grille
    "white":        "#FFFFFF",
}

# Palette categorielle : 8 couleurs harmonieuses (style editorial moderne)
SISTA_PALETTE = [
    "#13263D",  # 1. Navy SISTA
    "#D4AC0D",  # 2. Gold deep SISTA
    "#4A6FA5",  # 3. Steel blue
    "#C97B63",  # 4. Terracotta
    "#52796F",  # 5. Forest sage
    "#7D5A8C",  # 6. Muted plum
    "#E8A87C",  # 7. Peach
    "#A8B5C7",  # 8. Soft gray-blue
]

# Colormap sequentielle SISTA (beige clair -> navy profond)
SISTA_CMAP = LinearSegmentedColormap.from_list(
    "sista_seq",
    [
        (0.00, "#F4F7FA"),
        (0.20, "#C5D2E0"),
        (0.50, "#7A92B0"),
        (0.80, "#2D4A6E"),
        (1.00, "#13263D"),
    ],
)


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ----------------------------------------------------------------------
#  Filtres
# ----------------------------------------------------------------------

MAX_UNIVARIATE = 15
MAX_BIVARIATE = 8

EXCLUDED_TYPES = {"identifiant", "date", "vide"}

# Patterns de noms de variables a EXCLURE du rapport analytique destine au commanditaire.
# Ces variables sont des metadonnees de collecte / pilotage SISTA et n'apportent aucune
# valeur metier pour le rapport client.
EXCLUDED_NAME_PATTERNS = [
    # --- Metadonnees Kobo / ODK / XLSForm ---
    "_uuid", "_submission_time", "_submitted_by", "_validation_status",
    "_attachments", "_geolocation", "_tags", "_notes", "_index", "_parent_index",
    "_status", "_xform_id", "__version__", "formhub/", "meta/",
    "deviceid", "subscriberid", "simserial", "phonenumber", "username", "audit",
    "imei", "serial_number", "device_id",

    # --- Variables temporelles / pilotage de collecte ---
    "start", "end", "today", "starttime", "endtime",
    "duree_enquete", "duree_collecte", "duration", "duree_obs", "duree_obser",
    "date_collecte", "date_saisie", "heure_debut", "heure_fin",

    # --- Variables d'enqueteur / agent de collecte ---
    "enqueteur", "enquêteur", "enumerator", "interviewer", "agent_",
    "_agent", "surveyor", "investigator", "releveur", "operateur", "operator",
    "collecteur", "collector", "fieldworker", "field_worker", "data_collector",
    "nom_enq", "id_enq", "code_enq", "enum_name", "enum_id", "enum_code",
    "superviseur", "supervisor", "controleur", "controller",

    # --- Coordonnees GPS ---
    "gps", "latitude", "longitude", "altitude", "geopoint", "geoshape",
    "_lat", "_lon", "_lng", "_coord", "coordonnees",

    # --- Identifiants techniques internes ---
    "uuid", "guid", "session_id", "num_quest", "numero_quest",
    "num_obser", "numero_obs", "code_quest", "ref_quest",

    # --- Champs libres ou contacts (non analysables statistiquement) ---
    "commentaire", "comment", "comments", "remarque", "remark", "note",
    "observation_libre", "autre_precis", "autres_precise", "specifier",
    "telephone", "_phone", "tel_", "numero_tel", "mobile",
    "email", "courriel", "adresse_email",
    "nom_complet", "prenom", "nom_famille", "first_name", "last_name",
    "nom_repondant", "nom_chef",
]


def _is_analyzable(var: dict, mp: dict | None = None) -> bool:
    """Determine si une variable doit etre analysee dans le rapport client.

    Args:
        var : dict de profil de variable
        mp  : mapping des colonnes-cles {id, enqueteur, start, end, lat, lon}
              -> ces colonnes sont EXCLUES car techniques / de pilotage
    """
    name = var["name"]
    name_lower = name.lower()

    # ---- Exclusion DYNAMIQUE des colonnes-cles declarees par l'utilisateur ----
    # Le rapport est destine au COMMANDITAIRE de l'enquete (entreprise / ONG),
    # pas au bureau d'etudes SISTA. Les colonnes techniques de pilotage interne
    # (identifiant, enqueteur, horodatages, GPS) n'ont aucun interet metier.
    if mp:
        technical_cols = {
            mp.get(k) for k in ("id", "enqueteur", "start", "end", "lat", "lon")
            if mp.get(k)
        }
        if name in technical_cols:
            return False

    # ---- Exclusions par type ----
    if var["type"] in EXCLUDED_TYPES:
        return False

    # ---- Exclusions de qualite ----
    if var["fill_rate"] < 20:
        return False
    if var["uniques"] <= 1:
        return False

    # ---- Exclusions par pattern de nom (metadonnees de collecte) ----
    if any(p in name_lower for p in EXCLUDED_NAME_PATTERNS):
        return False
    if name.startswith("_") or name.startswith("__"):
        return False

    return True


def _get_exclusion_reason(var: dict, mp: dict | None = None) -> str | None:
    """Renvoie la raison d'exclusion d'une variable (None si elle est retenue).

    Permet a l'utilisateur de savoir POURQUOI une variable a ete ecartee.
    """
    name = var["name"]
    name_lower = name.lower()

    # 1. Colonne-cle declaree par l'utilisateur (priorite explicite)
    if mp:
        for k in ("id", "enqueteur", "start", "end", "lat", "lon"):
            if mp.get(k) == name:
                labels = {
                    "id": "Identifiant unique declare",
                    "enqueteur": "Colonne enqueteur declaree",
                    "start": "Horodatage de debut declare",
                    "end": "Horodatage de fin declare",
                    "lat": "Latitude GPS declaree",
                    "lon": "Longitude GPS declaree",
                }
                return labels.get(k, f"Colonne-cle declaree ({k})")

    # 2. Type non analysable
    if var["type"] in EXCLUDED_TYPES:
        labels = {
            "identifiant": "Type identifiant (non analysable)",
            "date":        "Type date (non pertinent pour le client)",
            "vide":        "Variable entierement vide",
        }
        return labels.get(var["type"], f"Type non analysable ({var['type']})")

    # 3. Qualite insuffisante
    if var["fill_rate"] < 20:
        return f"Taux de remplissage trop faible ({var['fill_rate']}%)"
    if var["uniques"] <= 1:
        return "Variable constante (1 seule valeur)"

    # 4. Pattern de nom (metadonnee de collecte)
    for p in EXCLUDED_NAME_PATTERNS:
        if p in name_lower:
            return f"Metadonnee technique (motif '{p}')"

    # 5. Prefixe _ (variable systeme)
    if name.startswith("_"):
        return "Variable systeme (prefixe '_')"

    return None  # Variable retenue


def compute_analysis_scope(profile: dict, mp: dict | None = None) -> dict:
    """Calcule le perimetre d'analyse : variables retenues vs exclues.

    Cette fonction expose AU CLIENT et AU RESPONSABLE QC ce qui sera
    effectivement analyse dans le rapport, et POURQUOI certaines variables
    ont ete ecartees. Cle pour la transparence du livrable.

    Args:
        profile : profil complet des variables (issu de profile_dataset)
        mp      : mapping des colonnes-cles {id, enqueteur, start, end, lat, lon}

    Returns:
        {
          "retained":     [{name, label, type, fill_rate, uniques}],
          "excluded":     [{name, label, type, reason}],
          "n_total":      nombre total de variables,
          "n_retained":   nombre retenu pour le rapport client,
          "n_excluded":   nombre exclu (techniques / metadonnees / qualite),
          "exclusion_summary": {reason: n}  -- regroupement par raison
        }
    """
    retained, excluded = [], []
    exclusion_summary = {}

    for v in profile.get("variables", []):
        reason = _get_exclusion_reason(v, mp)
        base = {
            "name":      v["name"],
            "label":     v.get("label", "") or v["name"],
            "type":      v["type"],
        }
        if reason is None:
            retained.append({
                **base,
                "fill_rate": v["fill_rate"],
                "uniques":   v["uniques"],
            })
        else:
            excluded.append({**base, "reason": reason})
            # Regroupement haut-niveau pour le resume
            high_level = reason.split(" (")[0].split(" '")[0]
            exclusion_summary[high_level] = exclusion_summary.get(high_level, 0) + 1

    return {
        "retained":          retained,
        "excluded":          excluded,
        "n_total":           len(retained) + len(excluded),
        "n_retained":        len(retained),
        "n_excluded":        len(excluded),
        "exclusion_summary": exclusion_summary,
    }


# ----------------------------------------------------------------------
#  PASSE 1 : PLANIFICATEUR IA
# ----------------------------------------------------------------------

PLANNER_SYSTEM = """Tu es un expert statisticien specialise dans l'analyse d'enquetes.
Ton role : produire un PLAN D'ANALYSE pour un RAPPORT DESTINE AU COMMANDITAIRE
de l'enquete (entreprise, ONG, organisation publique, bailleur de fonds).

CE RAPPORT N'EST PAS UN RAPPORT INTERNE DE PILOTAGE TERRAIN.
Il doit faire ressortir des INSIGHTS METIER UTILES pour le commanditaire :
- caracteristiques sociodemographiques de la population enquetee
- comportements, opinions, conditions de vie, niveaux de vie
- variables d'interet specifiques au theme de l'enquete
- croisements pertinents qui revelent des segments ou des disparites
- patterns interessants pour la prise de decision

Tu identifies :
1. Les variables sociodemographiques (age, sexe, region, education, profession, revenu, statut...)
2. Les variables principales d'interet (selon le theme de l'enquete)
3. Les croisements significatifs (sociodemo x variable principale, ou variable x variable)

REGLES STRICTES :
- IGNORE totalement les variables techniques / logistiques internes :
  identifiants (uuid, codes systeme), metadonnees de collecte (horodatage, GPS),
  informations sur les enqueteurs (nom, code, duree). Ces variables sont absentes
  de la liste fournie ci-dessous : ne tente JAMAIS d'en ajouter ou d'en inferer.
- Concentre-toi UNIQUEMENT sur les variables qui apportent une information
  utile au commanditaire pour comprendre sa population cible.

Tu reponds UNIQUEMENT en JSON valide, sans markdown, sans texte autour.
"""

PLANNER_USER_TEMPLATE = """CONTEXTE DE L'ENQUETE
Type : {survey_type}
Description : {survey_description}
Population : {survey_population}
Eligibilite : {survey_eligibility}

VARIABLES DISPONIBLES (deja filtrees) :
{variables_json}

CONSIGNES :
- Selectionne au MAXIMUM {max_uni} variables pour l'analyse univariee (les plus pertinentes)
- Propose au MAXIMUM {max_biv} croisements bivaries pertinents
- Classe les variables univariees en deux categories : "sociodemo" ou "principale"
- Pour chaque croisement, justifie en 1 phrase pourquoi il est pertinent

FORMAT DE REPONSE (JSON strict) :
{{
  "univariate": [
    {{"name": "nom_variable", "category": "sociodemo|principale", "priority": "high|medium"}}
  ],
  "bivariate": [
    {{"var1": "nom_var1", "var2": "nom_var2", "rationale": "pourquoi ce croisement est interessant"}}
  ]
}}
"""


def _build_planner_input(variables: list) -> list:
    out = []
    for v in variables:
        item = {
            "name": v["name"],
            "label": v.get("label", "")[:80],
            "type": v["type"],
            "uniques": v["uniques"],
            "fill_rate": v["fill_rate"],
            "examples": v.get("examples", [])[:3],
        }
        if v["type"] == "numerique" and v.get("stats"):
            item["stats"] = {
                "min": v["stats"].get("min"),
                "max": v["stats"].get("max"),
                "mean": v["stats"].get("mean"),
            }
        out.append(item)
    return out


def plan_report(api: str, api_key: str, profile: dict,
                survey_context: dict, progress_cb=None,
                mp: dict | None = None) -> dict:
    """Plan d'analyse IA. Filtre prealablement les variables techniques
    via _is_analyzable(var, mp) en utilisant le mapping de colonnes-cles
    declare par l'utilisateur."""
    if progress_cb:
        progress_cb("Filtrage des variables pertinentes pour le commanditaire...")

    analyzable = [v for v in profile["variables"] if _is_analyzable(v, mp)]
    if not analyzable:
        return {"univariate": [], "bivariate": []}

    if progress_cb:
        n_total = len(profile["variables"])
        n_filtered = n_total - len(analyzable)
        progress_cb(
            f"{len(analyzable)} variables retenues "
            f"({n_filtered} exclues car techniques / metadonnees). "
            f"L'IA prepare le plan..."
        )

    vars_json = json.dumps(_build_planner_input(analyzable), ensure_ascii=False, indent=2)
    user_prompt = PLANNER_USER_TEMPLATE.format(
        survey_type=survey_context.get("type", "Non specifie"),
        survey_description=survey_context.get("description", "Non specifie"),
        survey_population=survey_context.get("population", "Non specifie"),
        survey_eligibility=survey_context.get("eligibility", "Non specifie"),
        variables_json=vars_json,
        max_uni=MAX_UNIVARIATE,
        max_biv=MAX_BIVARIATE,
    )

    model = ai_agent.API_CONFIG[api]["model_smart"]
    resp = ai_agent._call_with_retry(api, api_key, model,
                                       PLANNER_SYSTEM, user_prompt, max_tokens=3000)
    plan = ai_agent._extract_json(resp["text"])

    plan["univariate"] = plan.get("univariate", [])[:MAX_UNIVARIATE]
    plan["bivariate"] = plan.get("bivariate", [])[:MAX_BIVARIATE]

    valid_names = {v["name"] for v in analyzable}
    plan["univariate"] = [u for u in plan["univariate"] if u.get("name") in valid_names]
    plan["bivariate"] = [
        b for b in plan["bivariate"]
        if b.get("var1") in valid_names and b.get("var2") in valid_names
    ]

    if progress_cb:
        progress_cb(f"Plan : {len(plan['univariate'])} univariees + {len(plan['bivariate'])} croisements")
    return plan


# ----------------------------------------------------------------------
#  CALCULS
# ----------------------------------------------------------------------

def compute_univariate(df: pd.DataFrame, var_info: dict) -> dict:
    col = var_info["name"]
    vtype = var_info["type"]
    series = df[col].dropna()

    result = {
        "name": col,
        "label": var_info.get("label", "") or col,
        "type": vtype,
        "n": len(series),
        "fill_rate": var_info["fill_rate"],
    }

    if vtype == "numerique":
        num = pd.to_numeric(
            series.astype(str).str.replace(",", ".", regex=False),
            errors="coerce"
        ).dropna()
        if len(num) == 0:
            result["error"] = "Pas de valeurs numeriques valides"
            return result
        result["stats"] = {
            "min": round(float(num.min()), 2),
            "max": round(float(num.max()), 2),
            "mean": round(float(num.mean()), 2),
            "median": round(float(num.median()), 2),
            "std": round(float(num.std()), 2) if len(num) > 1 else 0,
            "q1": round(float(num.quantile(0.25)), 2),
            "q3": round(float(num.quantile(0.75)), 2),
        }
        result["values_numeric"] = num.tolist()
        result["chart"] = "histogram_kde"
    else:
        vc = series.astype(str).value_counts()
        n_total = len(series)
        if len(vc) > 20:
            top = vc.head(10)
            autres = vc.iloc[10:].sum()
            vc_display = pd.concat([top, pd.Series({"Autres": autres})])
            result["truncated"] = True
        else:
            vc_display = vc
            result["truncated"] = False

        result["distribution"] = [
            {
                "modalite": str(k),
                "effectif": int(v),
                "pourcentage": round(100 * v / n_total, 1),
            }
            for k, v in vc_display.items()
        ]

        # Selection intelligente du type de graphe (v3 - diversification)
        # Critères : nb modalités + concentration (part du top1) + longueur libellés
        n_modal = len(vc_display)
        top1_pct = (vc_display.iloc[0] / n_total * 100) if n_total > 0 else 0
        max_label_len = max((len(str(k)) for k in vc_display.index), default=10)
        # Concentration "forte" si le top1 fait > 50% du total
        is_concentrated = top1_pct > 50

        if n_modal <= 2:
            # Binaire / quasi-binaire : donut reste le plus parlant
            result["chart"] = "donut"
        elif n_modal <= 5:
            # Peu de modalités : barres verticales = comparaison directe impactante
            # Sauf si une catégorie écrase tout : alors donut pour montrer le poids
            result["chart"] = "donut" if is_concentrated else "bar_vertical"
        elif n_modal <= 8:
            # 6-8 modalités : choix selon longueur des libellés
            #   labels longs -> horizontal (lisible)
            #   labels courts -> donut (esthétique) ou bar_vertical (comparaison)
            if max_label_len > 12:
                result["chart"] = "bar_horizontal"
            else:
                result["chart"] = "donut" if is_concentrated else "bar_vertical"
        elif n_modal <= 15:
            # 9-15 modalités : Pareto pour identifier les "vital few" 80/20
            # Pertinent quand la distribution est inegale (concentration moyenne+)
            if top1_pct > 25:
                result["chart"] = "pareto"
            else:
                result["chart"] = "bar_horizontal"
        else:
            # 16+ : lollipop reste le plus efficace pour beaucoup de modalités
            result["chart"] = "lollipop"
    return result


def compute_bivariate(df: pd.DataFrame, var1_info: dict, var2_info: dict) -> dict:
    col1, col2 = var1_info["name"], var2_info["name"]
    t1, t2 = var1_info["type"], var2_info["type"]

    result = {
        "var1": col1, "var2": col2,
        "label1": var1_info.get("label", "") or col1,
        "label2": var2_info.get("label", "") or col2,
        "type1": t1, "type2": t2,
    }

    sub = df[[col1, col2]].dropna()
    result["n"] = len(sub)

    if len(sub) == 0:
        result["error"] = "Aucune ligne avec les deux variables renseignees"
        return result

    # ---- CAT x CAT : heatmap OU stacked_100 selon la taille
    if t1 != "numerique" and t2 != "numerique":
        top1 = sub[col1].astype(str).value_counts().head(8).index
        top2 = sub[col2].astype(str).value_counts().head(8).index
        sub_f = sub[sub[col1].astype(str).isin(top1) & sub[col2].astype(str).isin(top2)]
        ct = pd.crosstab(
            sub_f[col1].astype(str), sub_f[col2].astype(str),
            normalize="index"
        ) * 100
        ct = ct.round(1)
        result["crosstab"] = {
            "index": ct.index.tolist(),
            "columns": ct.columns.tolist(),
            "values": ct.values.tolist(),
        }
        # Selection intelligente :
        #   - Petit tableau (peu de modalites par dimension) -> stacked_100 (lisible)
        #   - Grand tableau -> heatmap (synthese visuelle)
        n_idx = len(ct.index)
        n_cols = len(ct.columns)
        if n_idx <= 6 and n_cols <= 5:
            result["chart"] = "stacked_100"
        else:
            result["chart"] = "heatmap"

    # ---- CAT x NUM : boxplot
    elif (t1 == "numerique") != (t2 == "numerique"):
        num_col = col1 if t1 == "numerique" else col2
        cat_col = col2 if t1 == "numerique" else col1
        sub_calc = sub.copy()
        sub_calc[num_col] = pd.to_numeric(
            sub_calc[num_col].astype(str).str.replace(",", ".", regex=False),
            errors="coerce"
        )
        sub_calc = sub_calc.dropna(subset=[num_col])
        top_cat = sub_calc[cat_col].astype(str).value_counts().head(8).index
        sub_calc = sub_calc[sub_calc[cat_col].astype(str).isin(top_cat)]

        agg = sub_calc.groupby(sub_calc[cat_col].astype(str))[num_col].agg(
            ["mean", "median", "count", "std"]
        ).round(2)
        result["aggregation"] = {
            "num_var": num_col,
            "cat_var": cat_col,
            "rows": [
                {"modalite": idx, "moyenne": float(r["mean"]),
                 "mediane": float(r["median"]), "effectif": int(r["count"]),
                 "ecart_type": float(r["std"]) if pd.notna(r["std"]) else 0}
                for idx, r in agg.iterrows()
            ]
        }
        # Valeurs brutes par groupe pour le boxplot (capees a 500 par groupe)
        raw_groups = {"labels": [], "values": []}
        for cat in agg.index:
            vals = sub_calc[sub_calc[cat_col].astype(str) == cat][num_col].dropna()
            if len(vals) > 0:
                if len(vals) > 500:
                    vals = vals.sample(n=500, random_state=42)
                raw_groups["labels"].append(str(cat))
                raw_groups["values"].append(vals.tolist())
        result["raw_groups"] = raw_groups
        # Selection intelligente :
        #   - Petit n par groupe (<=200) -> violin (montre la forme de distrib)
        #   - Gros n -> boxplot (plus rapide a lire pour gros volumes)
        n_per_group = [len(v) for v in raw_groups["values"]]
        max_n = max(n_per_group) if n_per_group else 0
        min_n = min(n_per_group) if n_per_group else 0
        # Violon : il faut au moins 5 valeurs par groupe (sinon le KDE est mauvais)
        if max_n <= 200 and min_n >= 5:
            result["chart"] = "violin"
        else:
            result["chart"] = "boxplot"

    # ---- NUM x NUM : scatter + regression
    else:
        n1 = pd.to_numeric(sub[col1].astype(str).str.replace(",", ".", regex=False), errors="coerce")
        n2 = pd.to_numeric(sub[col2].astype(str).str.replace(",", ".", regex=False), errors="coerce")
        mask = n1.notna() & n2.notna()
        n1, n2 = n1[mask], n2[mask]
        if len(n1) < 2:
            result["error"] = "Pas assez de paires numeriques"
            return result
        corr = float(n1.corr(n2))
        result["correlation"] = round(corr, 3)
        result["scatter"] = {"x": n1.tolist()[:500], "y": n2.tolist()[:500]}
        result["chart"] = "scatter_reg"

    return result


# ----------------------------------------------------------------------
#  HELPERS POUR LES GRAPHES (KDE, regression)
# ----------------------------------------------------------------------

def _compute_kde(values, num_points=200, bandwidth=None):
    """KDE gaussien simple, 100% numpy, vectorise."""
    values = np.asarray(values, dtype=float)
    values = values[~np.isnan(values)]
    n = len(values)
    if n < 5:
        return None, None
    if bandwidth is None:
        std = float(np.std(values, ddof=1)) if n > 1 else 1.0
        bandwidth = 1.06 * max(std, 1e-6) * n ** (-1 / 5)
    x_min, x_max = float(values.min()), float(values.max())
    if x_min == x_max:
        return None, None
    pad = (x_max - x_min) * 0.05
    x = np.linspace(x_min - pad, x_max + pad, num_points)
    # Vectorise : (num_points, n)
    diff = (x[:, None] - values[None, :]) / bandwidth
    density = np.exp(-0.5 * diff ** 2).sum(axis=1)
    density /= (n * bandwidth * np.sqrt(2 * np.pi))
    return x, density


def _compute_regression(x, y, num_points=100):
    """Regression lineaire + bande de confiance a 95%."""
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    mask = ~(np.isnan(x) | np.isnan(y))
    x, y = x[mask], y[mask]
    n = len(x)
    if n < 3:
        return None
    try:
        slope, intercept = np.polyfit(x, y, 1)
    except Exception:
        return None
    x_pred = np.linspace(float(x.min()), float(x.max()), num_points)
    y_pred = slope * x_pred + intercept
    y_fit = slope * x + intercept
    residuals = y - y_fit
    s_err = float(np.sqrt(np.sum(residuals ** 2) / max(n - 2, 1)))
    x_mean = float(np.mean(x))
    sx2 = float(np.sum((x - x_mean) ** 2))
    if sx2 < 1e-10:
        return None
    se_pred = s_err * np.sqrt(1 / n + (x_pred - x_mean) ** 2 / sx2)
    t = 1.96  # 95% CI
    return {
        "x": x_pred,
        "y": y_pred,
        "ci_low": y_pred - t * se_pred,
        "ci_high": y_pred + t * se_pred,
        "slope": float(slope),
        "intercept": float(intercept),
    }


# ----------------------------------------------------------------------
#  GRAPHES MODERNES (PNG en memoire)
# ----------------------------------------------------------------------

def _setup_sista_style():
    plt.rcParams.update({
        "font.family": "DejaVu Sans",
        "font.size": 11,
        "axes.titlesize": 14,
        "axes.titleweight": "bold",
        "axes.titlecolor": SISTA["navy"],
        "axes.titlepad": 14,
        "axes.titlelocation": "left",
        "axes.labelsize": 10,
        "axes.labelcolor": SISTA["gray"],
        "axes.edgecolor": SISTA["gray_light"],
        "axes.linewidth": 0.8,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "xtick.color": SISTA["gray"],
        "ytick.color": SISTA["gray"],
        "xtick.labelsize": 9,
        "ytick.labelsize": 9,
        "figure.facecolor": SISTA["white"],
        "axes.facecolor": SISTA["white"],
        "grid.color": SISTA["gray_light"],
        "grid.linewidth": 0.6,
        "grid.alpha": 0.6,
        "legend.frameon": False,
        "legend.fontsize": 9,
    })


def _fig_to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=SISTA["white"])
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---- 1. HISTOGRAMME + KDE -------------------------------------------------

def chart_histogram_kde(uni: dict) -> bytes:
    """Histogramme moderne avec courbe KDE en superposition."""
    _setup_sista_style()
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    values = np.asarray(uni["values_numeric"], dtype=float)
    values = values[~np.isnan(values)]

    if len(values) < 2:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    # Histogramme rempli en navy
    n_bins = min(30, max(10, int(np.sqrt(len(values)))))
    ax.hist(
        values, bins=n_bins, color=SISTA["navy"],
        edgecolor="white", linewidth=1.0, alpha=0.85, density=True,
    )

    # Courbe KDE en gold par-dessus
    x_kde, y_kde = _compute_kde(values)
    if x_kde is not None:
        ax.fill_between(x_kde, y_kde, alpha=0.18, color=SISTA["gold"])
        ax.plot(x_kde, y_kde, color=SISTA["gold_deep"], linewidth=2.5,
                label="Densite estimee (KDE)")

    # Lignes de reference
    median = uni["stats"]["median"]
    mean = uni["stats"]["mean"]
    ax.axvline(median, color=SISTA["navy_soft"], linestyle="--", linewidth=1.5,
               alpha=0.85, label=f"Mediane : {median}")
    ax.axvline(mean, color=SISTA["gold_deep"], linestyle=":", linewidth=1.8,
               alpha=0.9, label=f"Moyenne : {mean}")

    ax.set_title(uni["label"][:70])
    ax.set_xlabel("Valeur")
    ax.set_ylabel("Densite")
    ax.legend(loc="best")
    ax.grid(True, axis="y", alpha=0.3)

    return _fig_to_png_bytes(fig)


# ---- 2. DONUT CHART -------------------------------------------------------

def chart_donut(uni: dict) -> bytes:
    """Donut chart moderne avec total au centre et legende detaillee."""
    _setup_sista_style()
    fig, ax = plt.subplots(figsize=(8, 5))

    dist = uni["distribution"]
    labels = [d["modalite"][:25] for d in dist]
    sizes = [d["effectif"] for d in dist]
    total = sum(sizes)
    if total == 0:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    colors = SISTA_PALETTE[:len(labels)]

    wedges, _texts, autotexts = ax.pie(
        sizes,
        labels=None,
        colors=colors,
        autopct=lambda p: f"{p:.1f}%" if p > 5 else "",
        startangle=90,
        counterclock=False,
        wedgeprops={"width": 0.40, "edgecolor": "white", "linewidth": 2.5},
        textprops={"fontsize": 10, "fontweight": "bold", "color": "white"},
        pctdistance=0.80,
    )

    # Total au centre du donut
    ax.text(0, 0.10, f"{total:,}".replace(",", " "),
            ha="center", va="center",
            fontsize=24, fontweight="bold", color=SISTA["navy"])
    ax.text(0, -0.12, "reponses", ha="center", va="center",
            fontsize=11, color=SISTA["gray"])

    # Legende a droite avec effectifs + %
    legend_labels = [
        f"{lbl}  -  {sz} ({100 * sz / total:.1f}%)"
        for lbl, sz in zip(labels, sizes)
    ]
    ax.legend(wedges, legend_labels, loc="center left",
              bbox_to_anchor=(1.05, 0.5), frameon=False, fontsize=10)

    ax.set_title(uni["label"][:70], pad=14)
    return _fig_to_png_bytes(fig)


# ---- 3. LOLLIPOP CHART ----------------------------------------------------

def chart_lollipop(uni: dict) -> bytes:
    """Lollipop chart horizontal - alternative moderne aux barres."""
    _setup_sista_style()
    dist = uni["distribution"]
    fig, ax = plt.subplots(figsize=(8.5, max(4.5, len(dist) * 0.45)))

    labels = [d["modalite"][:35] for d in dist]
    values = [d["effectif"] for d in dist]
    pcts = [d["pourcentage"] for d in dist]
    y_pos = np.arange(len(labels))

    if not values or max(values) == 0:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    # Lignes (sticks)
    ax.hlines(y=y_pos, xmin=0, xmax=values,
              color=SISTA["navy_soft"], linewidth=2, alpha=0.7)

    # Points (heads) - tous en navy
    ax.scatter(values, y_pos, s=140, color=SISTA["navy"],
               edgecolor="white", linewidth=2, zorder=3)

    # Point gold pour la modalite maximale
    max_idx = int(np.argmax(values))
    ax.scatter(values[max_idx], y_pos[max_idx], s=220, color=SISTA["gold"],
               edgecolor=SISTA["navy"], linewidth=2, zorder=4)

    # Annotations a droite des points
    x_max = max(values)
    for x, y, v, p in zip(values, y_pos, values, pcts):
        ax.text(x + x_max * 0.015, y, f"{v}  ({p}%)",
                va="center", fontsize=9, color=SISTA["navy"],
                fontweight="bold")

    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    ax.set_xlabel("Effectif")
    ax.set_xlim(0, x_max * 1.22)
    ax.set_title(uni["label"][:70], pad=14)
    ax.grid(True, axis="x", alpha=0.3)
    ax.spines["left"].set_visible(False)
    ax.tick_params(axis="y", length=0)

    return _fig_to_png_bytes(fig)


# ---- 4. HEATMAP -----------------------------------------------------------

def chart_heatmap(biv: dict) -> bytes:
    """Heatmap pour croisement categoriel x categoriel (% par ligne)."""
    _setup_sista_style()
    ct = biv["crosstab"]
    values = np.array(ct["values"], dtype=float)
    rows = ct["index"]
    cols = ct["columns"]

    if values.size == 0:
        fig, ax = plt.subplots(figsize=(6, 3))
        ax.text(0.5, 0.5, "Croisement vide", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    fig, ax = plt.subplots(figsize=(
        max(6, len(cols) * 1.1 + 2.5),
        max(4, len(rows) * 0.55 + 1.5),
    ))

    vmax = max(float(values.max()), 1.0)
    im = ax.imshow(values, cmap=SISTA_CMAP, aspect="auto", vmin=0, vmax=vmax)

    # Annotations dans les cellules
    threshold = vmax * 0.55
    for i in range(len(rows)):
        for j in range(len(cols)):
            v = values[i, j]
            color = "white" if v > threshold else SISTA["navy"]
            ax.text(j, i, f"{v:.1f}%", ha="center", va="center",
                    color=color, fontsize=10, fontweight="bold")

    # Encadrer la valeur max en gold
    max_pos = np.unravel_index(np.argmax(values), values.shape)
    ax.add_patch(Rectangle(
        (max_pos[1] - 0.5, max_pos[0] - 0.5), 1, 1,
        fill=False, edgecolor=SISTA["gold"], linewidth=3,
    ))

    ax.set_xticks(np.arange(len(cols)))
    ax.set_xticklabels([str(c)[:20] for c in cols],
                       rotation=30, ha="right")
    ax.set_yticks(np.arange(len(rows)))
    ax.set_yticklabels([str(r)[:25] for r in rows])

    ax.set_title(f"{biv['label1']} x {biv['label2']}"[:80], pad=14)

    # Colorbar discrete
    cbar = plt.colorbar(im, ax=ax, shrink=0.75, aspect=18, pad=0.02)
    cbar.set_label("% (par ligne)", fontsize=9, color=SISTA["gray"])
    cbar.outline.set_visible(False)
    cbar.ax.tick_params(labelsize=8, colors=SISTA["gray"])

    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(axis="both", length=0)
    ax.set_xticks(np.arange(-0.5, len(cols), 1), minor=True)
    ax.set_yticks(np.arange(-0.5, len(rows), 1), minor=True)
    ax.grid(which="minor", color="white", linewidth=2)
    ax.tick_params(which="minor", bottom=False, left=False)

    plt.tight_layout()
    return _fig_to_png_bytes(fig)


# ---- 5. BOXPLOT HORIZONTAL ------------------------------------------------

def chart_boxplot(biv: dict) -> bytes:
    """Boxplot horizontal pour croisement categoriel x numerique."""
    _setup_sista_style()
    raw = biv.get("raw_groups", {})
    labels = raw.get("labels", [])
    groups = raw.get("values", [])

    if not groups or all(len(g) == 0 for g in groups):
        # Fallback : lollipop des moyennes si pas de raw values
        return _chart_grouped_lollipop_fallback(biv)

    fig, ax = plt.subplots(figsize=(8.5, max(4.5, len(labels) * 0.55)))

    bp = ax.boxplot(
        groups, vert=False, patch_artist=True,
        labels=[str(l)[:30] for l in labels],
        widths=0.55,
        showfliers=True,
        flierprops=dict(marker="o", markerfacecolor=SISTA["gold_deep"],
                        markersize=4, markeredgecolor=SISTA["gold_deep"],
                        alpha=0.45),
        medianprops=dict(color=SISTA["gold"], linewidth=2.5),
        boxprops=dict(facecolor=SISTA["navy_soft"],
                      edgecolor=SISTA["navy"], linewidth=1.2, alpha=0.75),
        whiskerprops=dict(color=SISTA["navy"], linewidth=1.2),
        capprops=dict(color=SISTA["navy"], linewidth=1.2),
    )

    # Diamant gold pour la moyenne
    means = [float(np.mean(g)) if len(g) else 0.0 for g in groups]
    ax.scatter(means, np.arange(1, len(groups) + 1),
               marker="D", s=55, color=SISTA["gold"],
               edgecolor=SISTA["navy"], linewidth=1.2, zorder=5,
               label="Moyenne")

    num_label = biv['label1'] if biv['type1'] == 'numerique' else biv['label2']
    cat_label = biv['label2'] if biv['type1'] == 'numerique' else biv['label1']
    ax.set_xlabel(num_label)
    ax.set_title(f"Distribution de {num_label} selon {cat_label}"[:80], pad=14)
    ax.legend(loc="best")
    ax.grid(True, axis="x", alpha=0.3)
    ax.invert_yaxis()
    ax.spines["left"].set_visible(False)
    ax.tick_params(axis="y", length=0)

    return _fig_to_png_bytes(fig)


def _chart_grouped_lollipop_fallback(biv: dict) -> bytes:
    """Lollipop des moyennes (fallback si raw_groups absent)."""
    _setup_sista_style()
    agg = biv["aggregation"]
    fig, ax = plt.subplots(figsize=(8.5, max(4.5, len(agg["rows"]) * 0.45)))
    labels = [r["modalite"][:30] for r in agg["rows"]]
    means = [r["moyenne"] for r in agg["rows"]]
    y_pos = np.arange(len(labels))
    if not means:
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)
    x_max = max(means)
    ax.hlines(y=y_pos, xmin=0, xmax=means, color=SISTA["navy_soft"],
              linewidth=2, alpha=0.7)
    ax.scatter(means, y_pos, s=140, color=SISTA["navy"],
               edgecolor="white", linewidth=2, zorder=3)
    for x, y, v in zip(means, y_pos, means):
        ax.text(x + x_max * 0.015, y, f"{v:.1f}", va="center",
                fontsize=9, color=SISTA["navy"], fontweight="bold")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels)
    ax.invert_yaxis()
    num_label = biv['label1'] if biv['type1'] == 'numerique' else biv['label2']
    cat_label = biv['label2'] if biv['type1'] == 'numerique' else biv['label1']
    ax.set_xlabel(f"Moyenne de {num_label}")
    ax.set_title(f"{num_label} par {cat_label}"[:80], pad=14)
    ax.grid(True, axis="x", alpha=0.3)
    ax.spines["left"].set_visible(False)
    return _fig_to_png_bytes(fig)


# ---- 6. SCATTER + REGRESSION + IC 95% -------------------------------------

def chart_scatter_regression(biv: dict) -> bytes:
    """Scatter avec droite de regression et bande de confiance a 95%."""
    _setup_sista_style()
    sc = biv["scatter"]
    x = np.asarray(sc["x"], dtype=float)
    y = np.asarray(sc["y"], dtype=float)
    mask = ~(np.isnan(x) | np.isnan(y))
    x, y = x[mask], y[mask]

    fig, ax = plt.subplots(figsize=(7.5, 5.5))

    if len(x) < 2:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    # Scatter avec alpha adaptatif a la densite
    n = len(x)
    alpha = 0.65 if n < 100 else 0.45 if n < 300 else 0.30
    size = 40 if n < 100 else 28 if n < 300 else 18
    ax.scatter(x, y, color=SISTA["navy_soft"], alpha=alpha, s=size,
               edgecolor="white", linewidth=0.5, zorder=2)

    # Regression + bande IC 95%
    reg = _compute_regression(x, y)
    if reg is not None:
        ax.fill_between(reg["x"], reg["ci_low"], reg["ci_high"],
                        color=SISTA["gold"], alpha=0.22, zorder=3,
                        label="IC 95%")
        ax.plot(reg["x"], reg["y"], color=SISTA["gold_deep"],
                linewidth=2.5, zorder=4,
                label=f"Tendance (r = {biv.get('correlation', 0):.2f})")
        ax.legend(loc="best")

    ax.set_xlabel(biv["label1"])
    ax.set_ylabel(biv["label2"])
    ax.set_title(f"{biv['label1']} vs {biv['label2']}"[:80], pad=14)
    ax.grid(True, alpha=0.3)

    return _fig_to_png_bytes(fig)


# ---- DISPATCHER -----------------------------------------------------------

# ======================================================================
#  NOUVEAUX TYPES DE GRAPHIQUES (v3 - diversification)
# ======================================================================

# ---- 7. BAR VERTICAL ------------------------------------------------------

def chart_bar_vertical(uni: dict) -> bytes:
    """Barres verticales modernes - impactantes pour 3-8 modalites.

    Plus carre que le donut, mieux que le lollipop pour faire ressortir
    les comparaisons d'effectifs directs.
    """
    _setup_sista_style()
    fig, ax = plt.subplots(figsize=(8, 4.8))

    dist = uni["distribution"]
    if not dist:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    labels = [d["modalite"][:18] for d in dist]
    sizes = [d["effectif"] for d in dist]
    total = sum(sizes)
    pcts = [round(100 * s / total, 1) if total else 0 for s in sizes]

    # Mise en evidence de la modalite dominante en gold
    max_idx = sizes.index(max(sizes))
    colors = [SISTA["gold_deep"] if i == max_idx else SISTA["navy_soft"]
              for i in range(len(sizes))]

    x = np.arange(len(labels))
    bars = ax.bar(x, sizes, color=colors, edgecolor="white", linewidth=1.5,
                  width=0.7)

    # Etiquettes au-dessus des barres : effectif + %
    max_h = max(sizes) if sizes else 1
    for bar, n, p in zip(bars, sizes, pcts):
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2,
                h + max_h * 0.02,
                f"{n}\n({p}%)",
                ha="center", va="bottom",
                fontsize=10, fontweight="bold", color=SISTA["navy"])

    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=20 if max(len(l) for l in labels) > 10 else 0,
                       ha="right" if max(len(l) for l in labels) > 10 else "center",
                       fontsize=10)
    ax.set_ylabel("Effectif")
    ax.set_title(uni["label"][:70], pad=14)
    ax.set_ylim(0, max_h * 1.18)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(True, axis="y", alpha=0.3)

    return _fig_to_png_bytes(fig)


# ---- 8. BAR HORIZONTAL ----------------------------------------------------

def chart_bar_horizontal(uni: dict) -> bytes:
    """Barres horizontales classiques - ideales pour libelles longs.

    Alternative au lollipop avec un aspect plus solide / corporate.
    """
    _setup_sista_style()
    dist = uni["distribution"]
    n_items = len(dist)
    fig, ax = plt.subplots(figsize=(8, max(3.5, 0.42 * n_items + 1.5)))

    if not dist:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    # Tri DESC pour lecture naturelle (top en haut)
    dist_sorted = sorted(dist, key=lambda d: d["effectif"], reverse=True)
    labels = [d["modalite"][:35] for d in dist_sorted]
    sizes = [d["effectif"] for d in dist_sorted]
    pcts = [d["pourcentage"] for d in dist_sorted]
    total = sum(sizes)

    # Gradient navy -> navy_soft du plus grand au plus petit
    n = len(sizes)
    colors = []
    for i in range(n):
        # Top en navy fonce, queue en navy plus doux
        ratio = i / max(1, n - 1)
        colors.append(SISTA["navy"] if ratio < 0.33
                      else SISTA["navy_soft"] if ratio < 0.66
                      else "#A8B5C7")

    y = np.arange(n)[::-1]  # inverse pour avoir le top en haut
    bars = ax.barh(y, sizes, color=colors, edgecolor="white", linewidth=1.2,
                   height=0.75)

    # Etiquettes en bout de barre : effectif (pct%)
    max_w = max(sizes) if sizes else 1
    for bar, n_val, p in zip(bars, sizes, pcts):
        w = bar.get_width()
        ax.text(w + max_w * 0.015, bar.get_y() + bar.get_height() / 2,
                f"{n_val} ({p}%)",
                ha="left", va="center",
                fontsize=9, fontweight="bold", color=SISTA["navy"])

    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=10)
    ax.set_xlabel("Effectif")
    ax.set_title(uni["label"][:70], pad=12)
    ax.set_xlim(0, max_w * 1.20)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.grid(True, axis="x", alpha=0.3)

    return _fig_to_png_bytes(fig)


# ---- 9. PARETO CHART ------------------------------------------------------

def chart_pareto(uni: dict) -> bytes:
    """Pareto : barres triees DESC + courbe cumulee.

    Met en evidence le principe 80/20 : combien de modalites expliquent
    quel pourcentage du total. Tres parlant pour reperer les dominantes.
    """
    _setup_sista_style()
    fig, ax1 = plt.subplots(figsize=(8.5, 5))

    dist = uni["distribution"]
    if not dist:
        ax1.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                 color=SISTA["gray"])
        ax1.set_axis_off()
        return _fig_to_png_bytes(fig)

    # Tri decroissant
    dist_sorted = sorted(dist, key=lambda d: d["effectif"], reverse=True)
    labels = [d["modalite"][:18] for d in dist_sorted]
    sizes = [d["effectif"] for d in dist_sorted]
    total = sum(sizes)
    cum_pct = np.cumsum(sizes) / total * 100 if total else np.zeros(len(sizes))

    x = np.arange(len(labels))

    # Barres en navy
    bars = ax1.bar(x, sizes, color=SISTA["navy_soft"], edgecolor="white",
                   linewidth=1.2, width=0.75, label="Effectif")
    # Coloration speciale du top 3
    for i, b in enumerate(bars[:3]):
        b.set_color(SISTA["navy"])
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels, rotation=30, ha="right", fontsize=9)
    ax1.set_ylabel("Effectif", color=SISTA["navy"])
    ax1.tick_params(axis="y", labelcolor=SISTA["navy"])
    ax1.spines["top"].set_visible(False)
    ax1.grid(True, axis="y", alpha=0.3)

    # Courbe cumulee sur axe secondaire en gold
    ax2 = ax1.twinx()
    ax2.plot(x, cum_pct, color=SISTA["gold_deep"], linewidth=2.5,
             marker="o", markersize=7, markerfacecolor=SISTA["gold"],
             markeredgecolor=SISTA["gold_deep"], markeredgewidth=1.5,
             label="% cumulé")
    # Ligne de reference 80%
    ax2.axhline(80, color="#C0392B", linestyle="--", linewidth=1.2, alpha=0.7)
    ax2.text(len(x) - 0.5, 82, "80%", color="#C0392B", fontsize=9,
             fontweight="bold", ha="right")

    ax2.set_ylabel("% cumulé", color=SISTA["gold_deep"])
    ax2.tick_params(axis="y", labelcolor=SISTA["gold_deep"])
    ax2.set_ylim(0, 105)
    ax2.spines["top"].set_visible(False)

    # Identifier les "vital few" (modalites avant 80%)
    n_vital = int(np.argmax(cum_pct >= 80)) + 1 if any(cum_pct >= 80) else len(labels)
    ax1.set_title(
        f"{uni['label'][:60]}  —  {n_vital} modalité(s) explique(nt) ~80% du total",
        pad=12, fontsize=11
    )

    return _fig_to_png_bytes(fig)


# ---- 10. STACKED BAR 100% (bivarie cat x cat) ----------------------------

def chart_stacked_100(biv: dict) -> bytes:
    """Barres empilees 100% : alternative a la heatmap pour cat x cat.

    Plus lisible que la heatmap quand on a peu de modalites en abscisse.
    Chaque barre = 100% et montre la repartition d'une modalite par rapport
    a l'autre variable.
    """
    _setup_sista_style()
    fig, ax = plt.subplots(figsize=(8.5, 5))

    ct = biv.get("crosstab", {})
    rows = ct.get("index", [])
    cols = ct.get("columns", [])
    values = ct.get("values", [])  # deja normalises en %

    if not rows or not cols:
        ax.text(0.5, 0.5, "Donnees insuffisantes", ha="center", va="center",
                color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    values_arr = np.asarray(values, dtype=float)  # shape (n_rows, n_cols)
    n_rows, n_cols = values_arr.shape

    y = np.arange(n_rows)
    left = np.zeros(n_rows)
    colors = SISTA_PALETTE[:n_cols] if n_cols <= len(SISTA_PALETTE) \
             else (SISTA_PALETTE * ((n_cols // len(SISTA_PALETTE)) + 1))[:n_cols]

    for j in range(n_cols):
        seg = values_arr[:, j]
        bars = ax.barh(y, seg, left=left, height=0.7, color=colors[j],
                       edgecolor="white", linewidth=1.2,
                       label=str(cols[j])[:18])
        # Etiquette de pourcentage SI le segment est >= 8% (sinon illisible)
        for b, val, lf in zip(bars, seg, left):
            if val >= 8:
                ax.text(lf + val / 2, b.get_y() + b.get_height() / 2,
                        f"{val:.0f}%",
                        ha="center", va="center",
                        fontsize=9, color="white", fontweight="bold")
        left += seg

    ax.set_yticks(y)
    ax.set_yticklabels([str(r)[:25] for r in rows], fontsize=10)
    ax.set_xlabel(f"Répartition (%) de {biv.get('label2', biv.get('var2', ''))[:30]}")
    ax.set_xlim(0, 100)
    ax.set_title(
        f"{biv.get('label1', biv.get('var1', ''))[:35]} × "
        f"{biv.get('label2', biv.get('var2', ''))[:35]}",
        pad=12, fontsize=11,
    )
    ax.legend(loc="lower center", bbox_to_anchor=(0.5, -0.18),
              ncol=min(n_cols, 4), frameon=False, fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(True, axis="x", alpha=0.3)
    plt.tight_layout()

    return _fig_to_png_bytes(fig)


# ---- 11. VIOLIN PLOT (bivarie cat x num) ----------------------------------

def chart_violin(biv: dict) -> bytes:
    """Violons : alternative au boxplot, montre la distribution complete.

    Plus moderne et plus informatif quand on veut voir la forme de la
    distribution (multi-modalite, asymetrie...).
    """
    _setup_sista_style()
    fig, ax = plt.subplots(figsize=(8.5, 4.8))

    rg = biv.get("raw_groups", {})
    labels = rg.get("labels", [])
    groups = rg.get("values", [])

    # Filtrer les groupes vides
    valid = [(l, v) for l, v in zip(labels, groups) if v and len(v) >= 2]
    if not valid:
        ax.text(0.5, 0.5, "Donnees insuffisantes pour un violon",
                ha="center", va="center", color=SISTA["gray"])
        ax.set_axis_off()
        return _fig_to_png_bytes(fig)

    labels = [l[:22] for l, _ in valid]
    groups = [v for _, v in valid]

    parts = ax.violinplot(groups, showmeans=False, showmedians=False,
                          showextrema=False, widths=0.78)
    # Style des violons : navy_soft remplissage + bord navy
    for body in parts["bodies"]:
        body.set_facecolor(SISTA["navy_soft"])
        body.set_edgecolor(SISTA["navy"])
        body.set_linewidth(1.5)
        body.set_alpha(0.75)

    # Surperposer mediane (point gold) + boxplot fin par-dessus
    bp = ax.boxplot(groups, widths=0.15, showcaps=True, showfliers=False,
                    patch_artist=True,
                    boxprops=dict(facecolor=SISTA["navy"], edgecolor="white",
                                  linewidth=1.5),
                    medianprops=dict(color=SISTA["gold"], linewidth=2.5),
                    whiskerprops=dict(color=SISTA["navy"], linewidth=1.2),
                    capprops=dict(color=SISTA["navy"], linewidth=1.5))

    # Moyenne en losange gold par-dessus
    for i, vals in enumerate(groups, start=1):
        ax.scatter([i], [np.mean(vals)],
                   marker="D", color=SISTA["gold_deep"], s=55,
                   edgecolor="white", linewidth=1.5, zorder=4)

    ax.set_xticks(np.arange(1, len(labels) + 1))
    ax.set_xticklabels(labels, rotation=20 if max(len(l) for l in labels) > 10 else 0,
                       ha="right" if max(len(l) for l in labels) > 10 else "center",
                       fontsize=10)

    agg = biv.get("aggregation", {})
    num_var = agg.get("num_var", "")
    cat_var = agg.get("cat_var", "")
    ax.set_ylabel(num_var)
    ax.set_xlabel(cat_var)
    ax.set_title(
        f"Distribution de {num_var[:25]} selon {cat_var[:25]}",
        pad=12, fontsize=11,
    )

    # Petite legende explicative
    from matplotlib.patches import Patch
    from matplotlib.lines import Line2D
    legend_elems = [
        Patch(facecolor=SISTA["navy_soft"], edgecolor=SISTA["navy"],
              label="Distribution"),
        Line2D([0], [0], color=SISTA["gold"], lw=2.5, label="Médiane"),
        Line2D([0], [0], marker="D", color="w", markerfacecolor=SISTA["gold_deep"],
               markersize=8, markeredgecolor="white", label="Moyenne"),
    ]
    ax.legend(handles=legend_elems, loc="upper right", frameon=False, fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(True, axis="y", alpha=0.3)

    return _fig_to_png_bytes(fig)


# ======================================================================
#  Dispatcher unifie (avec nouveaux types)
# ======================================================================

def render_chart(item: dict, kind: str) -> bytes:
    """Dispatcher unifie des graphiques (avec fallbacks legacy)."""
    handlers = {
        # Types historiques (v1-v2)
        "histogram_kde": chart_histogram_kde,
        "donut":         chart_donut,
        "lollipop":      chart_lollipop,
        "heatmap":       chart_heatmap,
        "boxplot":       chart_boxplot,
        "scatter_reg":   chart_scatter_regression,
        # NOUVEAUX types (v3 - diversification)
        "bar_vertical":  chart_bar_vertical,
        "bar_horizontal":chart_bar_horizontal,
        "pareto":        chart_pareto,
        "stacked_100":   chart_stacked_100,
        "violin":        chart_violin,
        # Alias retro-compatibles (anciens noms)
        "histogram":   chart_histogram_kde,
        "pie":         chart_donut,
        "bar":         chart_lollipop,
        "stacked_bar": chart_heatmap,
        "grouped_bar": chart_boxplot,
        "scatter":     chart_scatter_regression,
    }
    fn = handlers.get(kind)
    if fn is None:
        return b""
    try:
        return fn(item)
    except Exception as e:
        print(f"Erreur rendu chart '{kind}' : {e}")
        return b""


# ----------------------------------------------------------------------
#  PASSE 2 : REDACTEUR IA
# ----------------------------------------------------------------------

WRITER_SYSTEM = """Tu es un analyste-redacteur expert en enquetes statistiques.
Ton role : interpreter des chiffres deja calcules et produire un texte clair, professionnel.

REGLES STRICTES :
- N'invente AUCUN chiffre : utilise seulement ceux fournis
- Style : phrases courtes, ton neutre et factuel, francais professionnel
- Pas de jargon technique inutile
- 2 a 4 phrases par interpretation, jamais plus
- Tu reponds en JSON valide UNIQUEMENT, sans markdown
"""

WRITER_USER_TEMPLATE = """CONTEXTE DE L'ENQUETE
Type : {survey_type}
Description : {survey_description}
Population : {survey_population}

DONNEES A INTERPRETER (deja calculees) :
{data_json}

CONSIGNE :
Pour chaque analyse fournie (univariate_i, bivariate_i), redige une interpretation
en 2-4 phrases qui :
- Decrit le constat principal en chiffres
- Apporte un eclairage qualitatif lie au contexte de l'enquete
- Reste neutre, factuel, professionnel

Redige aussi une SYNTHESE EXECUTIVE de 4-6 phrases qui resume les principaux enseignements.

FORMAT JSON STRICT :
{{
  "executive_summary": "Texte de la synthese executive en 4-6 phrases.",
  "interpretations": {{
    "univariate_0": "Interpretation de la premiere analyse univariee...",
    "univariate_1": "...",
    "bivariate_0": "Interpretation du premier croisement...",
    ...
  }}
}}
"""


def write_interpretations(api: str, api_key: str,
                          univariate_results: list,
                          bivariate_results: list,
                          survey_context: dict,
                          progress_cb=None) -> dict:
    if progress_cb:
        progress_cb("L'IA redige les interpretations...")

    payload = {"univariate": [], "bivariate": []}

    for i, u in enumerate(univariate_results):
        item = {
            "id": f"univariate_{i}",
            "variable": u["label"],
            "type": u["type"],
            "n": u["n"],
        }
        if u["type"] == "numerique" and "stats" in u:
            item["stats"] = u["stats"]
        elif "distribution" in u:
            item["distribution"] = u["distribution"][:10]
        payload["univariate"].append(item)

    for i, b in enumerate(bivariate_results):
        item = {
            "id": f"bivariate_{i}",
            "var1": b["label1"],
            "var2": b["label2"],
            "n": b["n"],
        }
        if "crosstab" in b:
            item["crosstab_pct"] = {
                "rows": b["crosstab"]["index"],
                "cols": b["crosstab"]["columns"],
                "values": b["crosstab"]["values"],
            }
        elif "aggregation" in b:
            item["moyennes_par_groupe"] = b["aggregation"]["rows"]
        elif "correlation" in b:
            item["correlation"] = b["correlation"]
        payload["bivariate"].append(item)

    data_json = json.dumps(payload, ensure_ascii=False, indent=2)
    user_prompt = WRITER_USER_TEMPLATE.format(
        survey_type=survey_context.get("type", "Non specifie"),
        survey_description=survey_context.get("description", "Non specifie"),
        survey_population=survey_context.get("population", "Non specifie"),
        data_json=data_json,
    )

    model = ai_agent.API_CONFIG[api]["model_smart"]
    resp = ai_agent._call_with_retry(api, api_key, model,
                                       WRITER_SYSTEM, user_prompt, max_tokens=4000)
    parsed = ai_agent._extract_json(resp["text"])
    if progress_cb:
        progress_cb("Interpretations recues.")
    return parsed


# ======================================================================
#  ETAPE 1 : BUILD_REPORT_CONTENT (pipeline complet -> contenu JSON)
# ======================================================================

def build_report_content(
    api: str,
    api_key: str,
    df: pd.DataFrame,
    profile: dict,
    survey_context: dict,
    filename: str,
    qc_results: list = None,
    qc_stats: dict = None,
    progress_cb=None,
    mp: dict | None = None,
) -> dict:
    """
    Pipeline complet (plan IA + calculs + graphes + interpretations IA)
    -> dict serialisable contenant TOUT le rapport.

    Args:
        mp : mapping des colonnes-cles {id, enqueteur, start, end, lat, lon}
             Les colonnes referencees ici sont EXCLUES de l'analyse car
             techniques (rapport destine au commanditaire, pas au bureau d'etudes).
    """
    if progress_cb:
        progress_cb("Demarrage de la generation du rapport...")

    # Passe 1 : plan (filtre les variables techniques via mp)
    plan = plan_report(api, api_key, profile, survey_context, progress_cb, mp=mp)

    if not plan.get("univariate") and not plan.get("bivariate"):
        raise RuntimeError("L'IA n'a propose aucune analyse exploitable.")

    var_by_name = {v["name"]: v for v in profile["variables"]}

    # Calculs + graphes univaries
    if progress_cb:
        progress_cb(f"Calcul de {len(plan['univariate'])} analyses univariees...")
    univariate_results = []
    for u in plan["univariate"]:
        if u["name"] in var_by_name:
            try:
                res = compute_univariate(df, var_by_name[u["name"]])
                res["category"] = u.get("category", "principale")
                if "error" not in res:
                    png_bytes = render_chart(res, res["chart"])
                    if png_bytes:
                        res["chart_base64"] = base64.b64encode(png_bytes).decode("ascii")
                univariate_results.append(res)
            except Exception as e:
                if progress_cb:
                    progress_cb(f"  ! Erreur sur {u['name']}: {e}")

    # Calculs + graphes bivaries
    if progress_cb:
        progress_cb(f"Calcul de {len(plan['bivariate'])} croisements...")
    bivariate_results = []
    for b in plan["bivariate"]:
        if b["var1"] in var_by_name and b["var2"] in var_by_name:
            try:
                res = compute_bivariate(df, var_by_name[b["var1"]], var_by_name[b["var2"]])
                res["rationale"] = b.get("rationale", "")
                if "error" not in res:
                    png_bytes = render_chart(res, res["chart"])
                    if png_bytes:
                        res["chart_base64"] = base64.b64encode(png_bytes).decode("ascii")
                bivariate_results.append(res)
            except Exception as e:
                if progress_cb:
                    progress_cb(f"  ! Erreur sur {b['var1']} x {b['var2']}: {e}")

    # Passe 2 : interpretations
    interpretations = write_interpretations(
        api, api_key, univariate_results, bivariate_results,
        survey_context, progress_cb
    )

    interps = interpretations.get("interpretations", {})
    for i, u in enumerate(univariate_results):
        u["interpretation"] = interps.get(f"univariate_{i}", "")
        # Retirer les donnees brutes lourdes (le graphe est deja rendu)
        u.pop("values_numeric", None)
    for i, b in enumerate(bivariate_results):
        b["interpretation"] = interps.get(f"bivariate_{i}", "")
        # Alleger le payload : on peut retirer les valeurs brutes par groupe
        b.pop("raw_groups", None)
        if "scatter" in b:
            b["scatter"] = {
                "x": b["scatter"]["x"][:200],
                "y": b["scatter"]["y"][:200],
            }

    # Methodologie auto
    summary = profile["summary"]
    methodology = (
        f"Le present rapport est base sur l'analyse de {summary['n_rows']} "
        f"observations collectees, mesurees sur {summary['n_vars']} variables. "
        f"Le jeu de donnees comprend {summary['n_numeric']} variables numeriques, "
        f"{summary['n_categorical']} variables categorielles, "
        f"{summary['n_date']} variables temporelles et "
        f"{summary['n_text']} variables textuelles. "
        f"Le taux global de remplissage est de {summary['global_fill_rate']}%."
    )

    # QC summary compact
    qc_summary = None
    if qc_results:
        qc_summary = [
            {
                "titre": r.get("titre", ""),
                "severite": r.get("severite", ""),
                "n_cas": r.get("n_cas", 0),
            }
            for r in qc_results if r.get("n_cas", 0) > 0
        ]

    # Annexe variables (compact)
    annex_variables = [
        {
            "name": v["name"],
            "type": v["type"],
            "fill_rate": v["fill_rate"],
            "uniques": v["uniques"],
        }
        for v in profile["variables"]
    ]

    # Perimetre d'analyse : transparence pour le client.
    # Indique CE QUI a ete analyse et CE QUI a ete ecarte (et pourquoi).
    analysis_scope = compute_analysis_scope(profile, mp)

    content = {
        "meta": {
            "filename": filename,
            "n_rows": summary["n_rows"],
            "n_vars": summary["n_vars"],
            "generated_at": datetime.now().isoformat(),
            "survey_context": survey_context,
        },
        "executive_summary": interpretations.get("executive_summary", ""),
        "methodology": methodology,
        "analysis_scope": analysis_scope,
        "univariate": univariate_results,
        "bivariate": bivariate_results,
        "qc_summary": qc_summary,
        "annex_variables": annex_variables,
    }

    if progress_cb:
        progress_cb("Contenu du rapport pret !")
    return content


# ======================================================================
#  ETAPE 2 : COMPOSE_WORD_FROM_CONTENT (contenu -> .docx)
# ======================================================================

LOGO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "assets", "sista_logo.png"
)


def _set_cell_background(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _rgb(SISTA["navy"])
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p


def _add_paragraph(doc: Document, text: str, italic=False, bold=False,
                   size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = _rgb(color)
    if align:
        p.alignment = align
    return p


def _add_image_from_base64(doc: Document, b64: str, width_inches=5.5):
    img_bytes = base64.b64decode(b64)
    stream = io.BytesIO(img_bytes)
    doc.add_picture(stream, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_distribution_table(doc: Document, distribution: list):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(["Modalite", "Effectif", "Pourcentage"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _rgb(SISTA["white"])
        run.font.size = Pt(11)
        _set_cell_background(hdr[i], SISTA["navy"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in distribution:
        cells = table.add_row().cells
        cells[0].text = str(row["modalite"])[:50]
        cells[1].text = str(row["effectif"])
        cells[2].text = f"{row['pourcentage']}%"
        for c in cells[1:]:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_stats_table(doc: Document, stats: dict):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Statistique", "Valeur"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _rgb(SISTA["white"])
        _set_cell_background(hdr[i], SISTA["navy"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    labels = {
        "min": "Minimum", "max": "Maximum", "mean": "Moyenne",
        "median": "Mediane", "std": "Ecart-type",
        "q1": "1er quartile (Q1)", "q3": "3e quartile (Q3)",
    }
    for key, lab in labels.items():
        if key in stats:
            cells = table.add_row().cells
            cells[0].text = lab
            cells[1].text = str(stats[key])
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_cover_page(doc: Document, survey_context: dict, filename: str, n_rows: int):
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.8))

    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("RAPPORT ANALYTIQUE")
    run_t.font.size = Pt(32)
    run_t.font.bold = True
    run_t.font.color.rgb = _rgb(SISTA["navy"])

    survey_type = survey_context.get("type") or "Enquete"
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(survey_type)
    run_s.font.size = Pt(18)
    run_s.font.color.rgb = _rgb(SISTA["gold_deep"])
    run_s.italic = True

    for _ in range(4):
        doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("Fichier source", filename),
        ("Nombre de repondants", f"{n_rows}"),
        ("Date de generation", datetime.now().strftime("%d/%m/%Y")),
    ]
    for i, (k, v) in enumerate(info_rows):
        c0, c1 = info_table.rows[i].cells
        c0.text = ""
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(k + " : ")
        run0.bold = True
        run0.font.color.rgb = _rgb(SISTA["navy"])
        c1.text = ""
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(v)
        run1.font.color.rgb = _rgb(SISTA["navy_deep"])

    for _ in range(6):
        doc.add_paragraph()

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_foot.add_run("SISTA Consult Mauritanie  -  © " + str(datetime.now().year))
    run_f.font.size = Pt(10)
    run_f.font.color.rgb = _rgb(SISTA["gold_deep"])
    run_f.bold = True

    doc.add_page_break()


def _render_uni_section(doc, uni):
    _add_heading(doc, uni["label"][:80], level=2)
    _add_paragraph(
        doc,
        f"Variable : {uni['name']} | Type : {uni['type']} | "
        f"Reponses : {uni['n']} ({uni['fill_rate']}% de remplissage)",
        italic=True, color=SISTA["gray"], size=9
    )

    if "error" in uni:
        _add_paragraph(doc, uni["error"], italic=True, color=SISTA["gray"])
        return

    if "chart_base64" in uni:
        try:
            _add_image_from_base64(doc, uni["chart_base64"], width_inches=5.5)
        except Exception as e:
            _add_paragraph(doc, f"(Graphe non disponible : {e})",
                           italic=True, color=SISTA["gray"])

    if uni["type"] == "numerique" and "stats" in uni:
        _add_stats_table(doc, uni["stats"])
    elif "distribution" in uni:
        _add_distribution_table(doc, uni["distribution"][:15])
        if uni.get("truncated"):
            _add_paragraph(doc, "(Affichage limite aux 10 modalites + autres)",
                           italic=True, size=9, color=SISTA["gray"])

    if uni.get("interpretation"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run_lab = p.add_run("Interpretation : ")
        run_lab.bold = True
        run_lab.font.color.rgb = _rgb(SISTA["gold_deep"])
        run_txt = p.add_run(uni["interpretation"])
        run_txt.font.color.rgb = _rgb(SISTA["navy_deep"])
    doc.add_paragraph()


def _render_biv_section(doc, biv):
    title = f"{biv['label1']} x {biv['label2']}"
    _add_heading(doc, title[:80], level=2)

    if biv.get("rationale"):
        _add_paragraph(doc, biv["rationale"], italic=True,
                       color=SISTA["gray"], size=10)

    if "error" in biv:
        _add_paragraph(doc, biv["error"], italic=True, color=SISTA["gray"])
        return

    if "chart_base64" in biv:
        try:
            _add_image_from_base64(doc, biv["chart_base64"], width_inches=6.0)
        except Exception as e:
            _add_paragraph(doc, f"(Graphe non disponible : {e})",
                           italic=True, color=SISTA["gray"])

    if biv.get("interpretation"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run_lab = p.add_run("Interpretation : ")
        run_lab.bold = True
        run_lab.font.color.rgb = _rgb(SISTA["gold_deep"])
        run_txt = p.add_run(biv["interpretation"])
        run_txt.font.color.rgb = _rgb(SISTA["navy_deep"])
    doc.add_paragraph()


def _add_qc_section_from_content(doc, qc_summary):
    _add_heading(doc, "Qualite des donnees", level=1)
    _add_paragraph(
        doc,
        "Cette section presente le bilan du controle qualite automatise. "
        "Les anomalies sont classees par gravite.",
        italic=True, color=SISTA["gray"]
    )

    if not qc_summary:
        _add_paragraph(doc, "Aucune anomalie detectee.", italic=True)
        return

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Test", "Gravite", "Cas detectes"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _rgb(SISTA["white"])
        _set_cell_background(hdr[i], SISTA["navy"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sev_label = {"high": "Eleve", "med": "Modere", "low": "Faible"}
    total = 0
    for r in qc_summary:
        cells = table.add_row().cells
        cells[0].text = r.get("titre", "")[:60]
        cells[1].text = sev_label.get(r.get("severite", ""), "")
        cells[2].text = str(r.get("n_cas", 0))
        total += r.get("n_cas", 0)
        for c in cells[1:]:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_paragraph(doc, f"Total des anomalies : {total}",
                   bold=True, color=SISTA["navy"])


def compose_word_from_content(content: dict) -> bytes:
    """
    Compose le .docx a partir d'un report_content (potentiellement edite).
    Retourne les bytes.
    """
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    meta = content["meta"]

    # Page de garde
    _add_cover_page(doc, meta["survey_context"], meta["filename"], meta["n_rows"])

    # Synthese executive
    _add_heading(doc, "Synthese executive", level=1)
    _add_paragraph(doc, content.get("executive_summary") or
                   "Ce rapport presente l'analyse descriptive de l'enquete.", size=11)
    doc.add_paragraph()

    # Methodologie
    _add_heading(doc, "Methodologie", level=1)
    _add_paragraph(doc, content.get("methodology", ""))
    survey_ctx = meta.get("survey_context", {})
    if survey_ctx.get("description"):
        doc.add_paragraph()
        _add_paragraph(doc, "Contexte de l'enquete :", bold=True)
        _add_paragraph(doc, survey_ctx["description"])
    if survey_ctx.get("population"):
        _add_paragraph(doc, "Population cible : " + survey_ctx["population"])
    doc.add_paragraph()

    # Perimetre d'analyse : transparence sur les variables retenues vs exclues
    scope = content.get("analysis_scope")
    if scope and scope.get("n_total"):
        _add_heading(doc, "Perimetre d'analyse", level=1)
        _add_paragraph(
            doc,
            f"Le fichier source contient {scope['n_total']} variables. "
            f"Pour ce rapport destine au commanditaire, {scope['n_retained']} variables "
            f"ont ete retenues pour l'analyse statistique. Les {scope['n_excluded']} variables "
            f"restantes ont ete ecartees car elles correspondent a des metadonnees techniques "
            f"de collecte (identifiants, horodatages, GPS, enqueteur, commentaires libres, "
            f"contacts personnels) sans valeur pour l'interpretation metier.",
            size=10,
        )

        # Resume par categorie d'exclusion
        if scope.get("exclusion_summary"):
            doc.add_paragraph()
            _add_paragraph(doc, "Resume des exclusions :", bold=True, size=10)
            sum_table = doc.add_table(rows=1, cols=2)
            sum_table.style = "Light Grid Accent 1"
            hdr = sum_table.rows[0].cells
            for i, h in enumerate(["Motif d'exclusion", "Nb variables"]):
                hdr[i].text = ""
                p = hdr[i].paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.color.rgb = _rgb(SISTA["white"])
                run.font.size = Pt(10)
                _set_cell_background(hdr[i], SISTA["navy"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for reason, n in sorted(scope["exclusion_summary"].items(),
                                     key=lambda x: -x[1]):
                cells = sum_table.add_row().cells
                cells[0].text = reason
                cells[1].text = str(n)
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for c in cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

        doc.add_paragraph()
    # Fin perimetre d'analyse

    # Univariees : sociodemo d'abord
    socio = [u for u in content["univariate"] if u.get("category") == "sociodemo"]
    principale = [u for u in content["univariate"] if u.get("category") != "sociodemo"]

    if socio:
        _add_heading(doc, "Profil des repondants", level=1)
        for u in socio:
            _render_uni_section(doc, u)

    if principale:
        _add_heading(doc, "Analyses descriptives", level=1)
        for u in principale:
            _render_uni_section(doc, u)

    # Croisements
    if content["bivariate"]:
        _add_heading(doc, "Analyses croisees", level=1)
        for b in content["bivariate"]:
            _render_biv_section(doc, b)

    # QC
    if content.get("qc_summary"):
        doc.add_page_break()
        _add_qc_section_from_content(doc, content["qc_summary"])

    # Annexe
    doc.add_page_break()
    _add_heading(doc, "Annexe : Liste des variables", level=1)
    annex_table = doc.add_table(rows=1, cols=4)
    annex_table.style = "Light Grid Accent 1"
    hdr = annex_table.rows[0].cells
    for i, h in enumerate(["Variable", "Type", "Remplissage", "Modalites"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = _rgb(SISTA["white"])
        run.font.size = Pt(10)
        _set_cell_background(hdr[i], SISTA["navy"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for v in content["annex_variables"]:
        cells = annex_table.add_row().cells
        cells[0].text = v["name"][:40]
        cells[1].text = v["type"]
        cells[2].text = f"{v['fill_rate']}%"
        cells[3].text = str(v["uniques"])
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()