"""
qc_report.py — Generation du rapport de Controle Qualite (Point 5 SISTA).

A l'issue de la phase de validation des anomalies par le responsable QC,
ce module compose un rapport Word formel qui sert de trace d'audit.

Distinction avec analytical_report.py :
  - analytical_report = analyse descriptive du contenu de l'enquete
                       (distributions, croisements, interpretation IA)
  - qc_report        = trace formelle des verifications qualite et des
                       decisions prises par le responsable QC

Structure du rapport :
  1. Page de garde (responsable QC, date, fichier source, ...)
  2. Synthese executive (compteurs + indicateur qualite global)
  3. Detail des regles QC basiques (avec statut de validation + commentaire)
  4. Detail des regles QC IA (avec statut de validation + commentaire)
  5. Bilan par enqueteur (anomalies confirmees attribuees a chacun)
  6. Recommandations / actions a mener
  7. Page de signature
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ======================================================================
#  Constantes : statuts de validation + charte SISTA
# ======================================================================

# Vocabulaire des statuts de validation
VALIDATION_STATUSES = {
    "confirmed":      "Confirmée",
    "false_positive": "Faux positif",
    "corrected":      "Déjà corrigée",
    "pending":        "En attente",
}

# Couleur d'affichage par statut (utilisée dans les tableaux Word)
STATUS_COLOR = {
    "confirmed":      "#C0392B",   # rouge — anomalie réelle
    "false_positive": "#7F8C8D",   # gris — règle s'est trompée
    "corrected":      "#27AE60",   # vert — déjà résolue
    "pending":        "#F39C12",   # orange — pas encore tranché
}

# Charte SISTA (alignée sur analytical_report.py)
SISTA = {
    "navy":      "#13263D",
    "navy_deep": "#0D1B2C",
    "gold":      "#EFC71A",
    "gold_deep": "#D4AC0D",
    "gray":      "#6B7280",
    "gray_light":"#E5E7EB",
    "white":     "#FFFFFF",
    "cream":     "#FAF6EC",
}

# Logo SISTA (meme chemin qu'analytical_report.py)
LOGO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "assets", "sista_logo.png"
)


# ======================================================================
#  API publique : identifiants et synthese
# ======================================================================

def make_item_id(kind: str, key: Any) -> str:
    """Construit l'identifiant unique d'une regle pour le stockage des validations.

    Exemples :
      - make_item_id('basic', 'doublons_lignes') -> 'basic:doublons_lignes'
      - make_item_id('ai', 0)                    -> 'ai:0'
    """
    return f"{kind}:{key}"


def get_validation(validations: dict, item_id: str) -> dict:
    """Recupere la validation d'un item ou renvoie un statut 'pending' par defaut."""
    if not validations:
        return {"status": "pending", "comment": ""}
    v = validations.get(item_id)
    if not v:
        return {"status": "pending", "comment": ""}
    return {
        "status": v.get("status", "pending"),
        "comment": v.get("comment", "") or "",
    }


def summarize_validations(
    validations: dict,
    qc_results: list,
    ai_rules: list | None = None,
    ai_result: dict | None = None,
    n_observations: int = 0,
) -> dict:
    """Calcule les compteurs globaux du rapport QC.

    Returns:
        {
          n_observations, n_rules_total, n_rules_basic, n_rules_ai,
          n_anomalies_total,
          by_status: {confirmed: N, false_positive: N, corrected: N, pending: N},
          taux_qualite : pourcentage d'observations sans anomalie confirmee
        }
    """
    rules_basic = [r for r in (qc_results or []) if r.get("severite") != "ok"]
    rules_ai = list(ai_rules or [])

    by_status = {k: 0 for k in VALIDATION_STATUSES}

    # Cas QC basique : 1 ligne par regle => statut s'applique a tous ses cas
    for r in rules_basic:
        item_id = make_item_id("basic", r.get("id") or r.get("titre", ""))
        v = get_validation(validations, item_id)
        by_status[v["status"]] += r.get("n_cas", 0)

    # Cas IA : on a une regle agregee par index, n_cas dispatches via cas_par_regle
    if rules_ai and ai_result:
        cas_par_regle = ai_result.get("cas_par_regle", {})
        for i in range(len(rules_ai)):
            item_id = make_item_id("ai", i)
            v = get_validation(validations, item_id)
            n_cas = int(cas_par_regle.get(str(i), cas_par_regle.get(i, 0)) or 0)
            by_status[v["status"]] += n_cas

    n_anomalies_total = sum(by_status.values())

    # Taux de qualite : on l'estime comme la part d'observations sans anomalie
    # confirmee (approche conservative — une observation peut etre touchee par
    # plusieurs anomalies, mais on n'a pas le compte exact des observations
    # uniques en cause au niveau du rapport).
    n_confirmed = by_status["confirmed"]
    if n_observations > 0:
        # Cap pour ne pas avoir un taux negatif si plusieurs anomalies / observation
        taux = max(0.0, (n_observations - n_confirmed) / n_observations * 100)
        taux_qualite = round(taux, 1)
    else:
        taux_qualite = 0.0

    return {
        "n_observations":    n_observations,
        "n_rules_total":     len(rules_basic) + len(rules_ai),
        "n_rules_basic":     len(rules_basic),
        "n_rules_ai":        len(rules_ai),
        "n_anomalies_total": n_anomalies_total,
        "by_status":         by_status,
        "taux_qualite":      taux_qualite,
    }


# ======================================================================
#  Helpers Word (autonomes pour eviter d'importer analytical_report)
# ======================================================================

def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_background(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold=False, italic=False, color=None, align=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = _rgb(color)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _rgb(SISTA["navy"])
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p


def _add_paragraph(doc: Document, text: str, italic=False, bold=False,
                   size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = _rgb(color)
    if align is not None:
        p.alignment = align
    return p


def _header_row(table, headers):
    """Formate la 1re ligne d'un tableau (fond navy + texte blanc)."""
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, color=SISTA["white"],
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _set_cell_background(hdr[i], SISTA["navy"])


# ======================================================================
#  SECTIONS DU RAPPORT
# ======================================================================

def _add_cover_page(doc: Document, metadata: dict, filename: str,
                    n_observations: int, survey_type: str = ""):
    """Page de garde du rapport QC."""
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.8))

    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("RAPPORT DE CONTRÔLE QUALITÉ")
    run_t.font.size = Pt(28)
    run_t.font.bold = True
    run_t.font.color.rgb = _rgb(SISTA["navy"])

    if survey_type:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_s = p_sub.add_run(survey_type)
        run_s.font.size = Pt(16)
        run_s.font.color.rgb = _rgb(SISTA["gold_deep"])
        run_s.italic = True

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s2 = p_sub2.add_run("Trace formelle de la validation des anomalies")
    run_s2.font.size = Pt(12)
    run_s2.font.color.rgb = _rgb(SISTA["gray"])
    run_s2.italic = True

    for _ in range(4):
        doc.add_paragraph()

    # Tableau d'identification
    info_rows = [
        ("Fichier source",       filename),
        ("Nombre d'observations", f"{n_observations}"),
        ("Date du contrôle",     metadata.get("date_validation") or datetime.now().strftime("%d/%m/%Y")),
        ("Responsable QC",       metadata.get("responsable_qc") or "—"),
    ]
    if metadata.get("organisation"):
        info_rows.append(("Organisation", metadata["organisation"]))

    info_table = doc.add_table(rows=len(info_rows), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info_rows):
        c0, c1 = info_table.rows[i].cells
        _set_cell_text(c0, k + " : ", bold=True, color=SISTA["navy"], size=11)
        _set_cell_text(c1, v, color=SISTA["navy_deep"], size=11)

    for _ in range(6):
        doc.add_paragraph()

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_foot.add_run(f"SISTA Consult Mauritanie  -  © {datetime.now().year}")
    run_f.font.size = Pt(10)
    run_f.font.color.rgb = _rgb(SISTA["gold_deep"])
    run_f.bold = True

    doc.add_page_break()


def _add_executive_summary(doc: Document, summary: dict):
    """Section 1 : synthese executive avec compteurs et taux de qualite."""
    _add_heading(doc, "1. Synthèse exécutive", level=1)

    _add_paragraph(
        doc,
        "Cette section présente la vue d'ensemble des contrôles effectués "
        "et de leur traitement par le responsable qualité.",
        italic=True, color=SISTA["gray"], size=10,
    )
    doc.add_paragraph()

    # Tableau de compteurs
    rows = [
        ("Observations analysées",    f"{summary['n_observations']}"),
        ("Règles exécutées (total)",  f"{summary['n_rules_total']} "
                                       f"({summary['n_rules_basic']} basiques + "
                                       f"{summary['n_rules_ai']} générées par IA)"),
        ("Anomalies détectées",        f"{summary['n_anomalies_total']}"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        _set_cell_text(c0, k, bold=True, color=SISTA["navy"], size=11)
        _set_cell_text(c1, v, color=SISTA["navy_deep"], size=11,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()
    _add_paragraph(doc, "Répartition des anomalies par décision",
                   bold=True, size=12, color=SISTA["navy"])

    # Tableau des statuts
    by_status = summary["by_status"]
    n_total = summary["n_anomalies_total"] or 1
    status_table = doc.add_table(rows=1, cols=3)
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_table.style = "Light Grid Accent 1"
    _header_row(status_table, ["Décision", "Nombre", "Pourcentage"])

    for status_key in ["confirmed", "false_positive", "corrected", "pending"]:
        n = by_status[status_key]
        if n == 0 and status_key == "pending":
            # Si rien en pending, on n'affiche pas la ligne (signe que tout est tranche)
            continue
        cells = status_table.add_row().cells
        _set_cell_text(cells[0], VALIDATION_STATUSES[status_key],
                       bold=True, color=STATUS_COLOR[status_key], size=11)
        _set_cell_text(cells[1], str(n), align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        pct = round(n / n_total * 100, 1) if n_total else 0
        _set_cell_text(cells[2], f"{pct}%", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    doc.add_paragraph()

    # Indicateur de qualite global
    taux = summary["taux_qualite"]
    if taux >= 90:
        verdict = "Qualité globale satisfaisante"
        verdict_color = "#27AE60"
    elif taux >= 70:
        verdict = "Qualité globale à améliorer"
        verdict_color = "#F39C12"
    else:
        verdict = "Qualité globale insuffisante — actions correctives requises"
        verdict_color = "#C0392B"

    _add_paragraph(doc, "Indicateur de qualité global",
                   bold=True, size=12, color=SISTA["navy"])
    _add_paragraph(
        doc,
        f"{taux}%  —  {verdict}",
        bold=True, size=14, color=verdict_color,
    )
    _add_paragraph(
        doc,
        f"Calcul : (observations sans anomalie confirmée) / (observations totales). "
        f"Une observation pouvant porter plusieurs anomalies, ce taux est une "
        f"estimation conservative.",
        italic=True, color=SISTA["gray"], size=9,
    )

    doc.add_page_break()


def _render_rules_section(doc: Document, heading: str, rules_data: list,
                          intro: str = ""):
    """Affiche un tableau de regles avec leur statut de validation.

    rules_data : liste de dicts avec
        { titre, severite, n_cas, status, status_label, comment, kind_label, regle_idx }
    """
    if not rules_data:
        _add_heading(doc, heading, level=1)
        _add_paragraph(doc, "Aucune règle dans cette catégorie.",
                       italic=True, color=SISTA["gray"], size=10)
        return

    _add_heading(doc, heading, level=1)
    if intro:
        _add_paragraph(doc, intro, italic=True, color=SISTA["gray"], size=10)
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    _header_row(table, ["#", "Règle / Test", "Cas", "Décision", "Commentaire"])

    # Largeur des colonnes (Word respecte plus ou moins)
    widths = [Cm(0.8), Cm(7.0), Cm(1.5), Cm(2.5), Cm(5.0)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    for i, r in enumerate(rules_data, start=1):
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _set_cell_text(cells[1], r["titre"], size=10, color=SISTA["navy_deep"])
        _set_cell_text(cells[2], str(r["n_cas"]),
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True)
        _set_cell_text(cells[3], r["status_label"],
                       bold=True, color=STATUS_COLOR.get(r["status"], SISTA["gray"]),
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        comment_text = r.get("comment") or "—"
        _set_cell_text(cells[4], comment_text,
                       italic=not r.get("comment"), size=9,
                       color=SISTA["gray"] if not r.get("comment") else SISTA["navy_deep"])

    doc.add_paragraph()


def _add_enqueteur_section(doc: Document, enq_summary: list,
                           validations: dict, qc_results: list,
                           ai_result: dict | None):
    """Section : bilan par enqueteur (anomalies CONFIRMEES uniquement).

    On recalcule en filtrant : seules les anomalies dont la regle est marquee
    'confirmed' sont attribuees a l'enqueteur.
    """
    _add_heading(doc, "4. Bilan par enquêteur", level=1)
    _add_paragraph(
        doc,
        "Anomalies confirmées attribuées à chaque enquêteur. "
        "Les faux positifs et corrections déjà effectuées sont exclus de ce bilan.",
        italic=True, color=SISTA["gray"], size=10,
    )
    doc.add_paragraph()

    # Re-aggregation : on compte uniquement les cas appartenant a une regle
    # avec statut 'confirmed'
    confirmed_basic_ids = set()
    for r in (qc_results or []):
        if r.get("severite") == "ok":
            continue
        item_id = make_item_id("basic", r.get("id") or r.get("titre", ""))
        v = get_validation(validations, item_id)
        if v["status"] == "confirmed":
            confirmed_basic_ids.add(r.get("id") or r.get("titre", ""))

    confirmed_ai_idx = set()
    if ai_result:
        # On lit aiRules implicitement via les indices presents dans validations
        for item_id, val in (validations or {}).items():
            if item_id.startswith("ai:") and val.get("status") == "confirmed":
                try:
                    confirmed_ai_idx.add(int(item_id.split(":", 1)[1]))
                except (ValueError, IndexError):
                    pass

    # Re-aggregate per enqueteur
    agg = {}

    def _add_to(enq, test_label):
        if not enq or enq == "—":
            return
        agg.setdefault(enq, {"nom": enq, "total": 0, "par_test": {}})
        agg[enq]["total"] += 1
        agg[enq]["par_test"][test_label] = agg[enq]["par_test"].get(test_label, 0) + 1

    # QC basique : on parcourt les lignes des regles confirmees
    for r in (qc_results or []):
        rid = r.get("id") or r.get("titre", "")
        if rid not in confirmed_basic_ids:
            continue
        for ligne in r.get("lignes", []):
            _add_to(ligne.get("_enqueteur"), r.get("titre", ""))

    # QC IA : on parcourt les cas, filtrés par _rule_idx confirmé
    if ai_result and confirmed_ai_idx:
        for cas in ai_result.get("lignes", []):
            ri = cas.get("_rule_idx")
            if ri in confirmed_ai_idx:
                _add_to(cas.get("Enqueteur") or cas.get("_enqueteur"),
                        f"Règle IA #{ri+1}")

    enq_list = sorted(agg.values(), key=lambda x: -x["total"])

    if not enq_list:
        _add_paragraph(
            doc,
            "Aucune anomalie confirmée attribuable à un enquêteur.",
            italic=True, color=SISTA["gray"], size=11,
        )
        doc.add_paragraph()
        return

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    _header_row(table, ["Enquêteur", "Anomalies confirmées", "Détail par règle"])

    for e in enq_list:
        cells = table.add_row().cells
        _set_cell_text(cells[0], e["nom"], bold=True,
                       color=SISTA["navy_deep"], size=10)
        # Code couleur seuil simple
        total = e["total"]
        couleur = "#C0392B" if total >= 10 else "#F39C12" if total >= 4 else "#27AE60"
        _set_cell_text(cells[1], str(total),
                       bold=True, color=couleur, size=11,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        details = " | ".join(f"{t} ({n})" for t, n in e["par_test"].items())
        _set_cell_text(cells[2], details, size=9, color=SISTA["gray"])

    doc.add_paragraph()


def _add_recommendations_section(doc: Document, rules_data: list,
                                  observations_generales: str = ""):
    """Section recommandations : agreger les actions des regles confirmees."""
    _add_heading(doc, "5. Recommandations et actions à mener", level=1)

    if observations_generales:
        _add_paragraph(doc, "Observations générales du responsable QC",
                       bold=True, color=SISTA["navy"], size=12)
        _add_paragraph(doc, observations_generales, size=11, color=SISTA["navy_deep"])
        doc.add_paragraph()

    confirmed = [r for r in rules_data if r.get("status") == "confirmed"]

    if not confirmed:
        _add_paragraph(
            doc,
            "Aucune anomalie n'a été confirmée. Aucune action corrective n'est requise.",
            italic=True, color=SISTA["gray"], size=11,
        )
        return

    _add_paragraph(
        doc,
        f"Les {len(confirmed)} règles ci-dessous ont été marquées comme "
        "réellement problématiques et appellent une action corrective.",
        italic=True, color=SISTA["gray"], size=10,
    )
    doc.add_paragraph()

    for i, r in enumerate(confirmed, start=1):
        _add_paragraph(
            doc,
            f"{i}. {r['titre']}",
            bold=True, color=SISTA["navy"], size=11,
        )
        if r.get("action"):
            _add_paragraph(
                doc,
                f"   Action proposée : {r['action']}",
                size=10, color=SISTA["navy_deep"],
            )
        if r.get("comment"):
            _add_paragraph(
                doc,
                f"   Commentaire du responsable QC : {r['comment']}",
                italic=True, size=10, color=SISTA["gray"],
            )
        _add_paragraph(
            doc,
            f"   Cas concernés : {r['n_cas']}",
            size=10, color=SISTA["gray"],
        )

    doc.add_paragraph()


def _add_signature_page(doc: Document, metadata: dict):
    """Page finale : signature du responsable QC."""
    doc.add_page_break()
    _add_heading(doc, "6. Validation et signature", level=1)
    _add_paragraph(
        doc,
        "Le présent rapport constitue la trace formelle des vérifications "
        "effectuées et des décisions prises sur la qualité des données.",
        size=11, color=SISTA["navy_deep"],
    )
    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau d'identification + bloc signature
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Nom du responsable QC",  metadata.get("responsable_qc") or "_____________________________"),
        ("Fonction / titre",       metadata.get("fonction") or "_____________________________"),
        ("Date",                   metadata.get("date_validation") or datetime.now().strftime("%d/%m/%Y")),
        ("Signature",              ""),
    ]
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        _set_cell_text(c0, k, bold=True, color=SISTA["navy"], size=11)
        _set_cell_text(c1, v, color=SISTA["navy_deep"], size=11)

    # Espace pour la signature
    for _ in range(4):
        doc.add_paragraph()
    _add_paragraph(
        doc,
        "____________________________________",
        align=WD_ALIGN_PARAGRAPH.CENTER, color=SISTA["gray"],
    )
    _add_paragraph(
        doc,
        "(Signature manuscrite)",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SISTA["gray"], size=9,
    )


# ======================================================================
#  Fonction principale : build_qc_report
# ======================================================================

def build_qc_report(
    filename: str,
    n_observations: int,
    qc_results: list,
    ai_rules: list | None,
    ai_result: dict | None,
    validations: dict,
    metadata: dict,
    survey_type: str = "",
) -> bytes:
    """Construit le rapport QC complet et renvoie les bytes du .docx.

    Args:
        filename             : nom du fichier source analysé
        n_observations       : nombre d'observations dans le fichier
        qc_results           : liste des résultats des tests QC basiques
        ai_rules             : liste des règles générées par l'IA (ou None)
        ai_result            : dict des résultats IA (avec 'lignes', 'cas_par_regle')
        validations          : dict { item_id: {status, comment} }
        metadata             : { responsable_qc, fonction, date_validation,
                                 organisation, observations_generales }
        survey_type          : type d'enquête (affiché en page de garde)

    Returns:
        bytes du fichier .docx
    """
    doc = Document()

    # --- Page de garde ---
    _add_cover_page(doc, metadata, filename, n_observations, survey_type)

    # --- Section 1 : synthese executive ---
    summary = summarize_validations(
        validations, qc_results, ai_rules, ai_result, n_observations,
    )
    _add_executive_summary(doc, summary)

    # --- Section 2 : detail des regles QC basiques ---
    basic_data = []
    for r in (qc_results or []):
        if r.get("severite") == "ok":
            continue
        item_id = make_item_id("basic", r.get("id") or r.get("titre", ""))
        v = get_validation(validations, item_id)
        basic_data.append({
            "titre":         r.get("titre", ""),
            "severite":      r.get("severite", "low"),
            "n_cas":         r.get("n_cas", 0),
            "status":        v["status"],
            "status_label":  VALIDATION_STATUSES.get(v["status"], v["status"]),
            "comment":       v["comment"],
            "action":        (r.get("explication") or {}).get("action", ""),
            "kind_label":    "QC basique",
        })

    _render_rules_section(
        doc,
        "2. Contrôles automatiques (QC basique)",
        basic_data,
        intro="Tests universels appliqués systématiquement à chaque fichier "
              "(doublons, valeurs manquantes, outliers, durées, GPS, etc.).",
    )

    # --- Section 3 : detail des regles QC IA ---
    ai_data = []
    if ai_rules and ai_result:
        cas_par_regle = ai_result.get("cas_par_regle", {})
        for i, rule in enumerate(ai_rules):
            item_id = make_item_id("ai", i)
            v = get_validation(validations, item_id)
            n_cas = int(cas_par_regle.get(str(i), cas_par_regle.get(i, 0)) or 0)
            ai_data.append({
                "titre":         rule.get("description", f"Règle {i+1}"),
                "severite":      "med",
                "n_cas":         n_cas,
                "status":        v["status"],
                "status_label":  VALIDATION_STATUSES.get(v["status"], v["status"]),
                "comment":       v["comment"],
                "action":        rule.get("action", ""),
                "kind_label":    "QC IA",
                "regle_idx":     i,
            })

    _render_rules_section(
        doc,
        "3. Règles générées par l'IA",
        ai_data,
        intro="Règles de cohérence métier produites par l'IA sur la base du "
              "dictionnaire, du contexte de l'enquête et du contenu du fichier.",
    )

    # --- Section 4 : bilan par enqueteur (filtré sur confirmées) ---
    _add_enqueteur_section(doc, [], validations, qc_results, ai_result)

    # --- Section 5 : recommandations ---
    all_rules_data = basic_data + ai_data
    _add_recommendations_section(
        doc, all_rules_data,
        observations_generales=metadata.get("observations_generales", ""),
    )

    # --- Section 6 : signature ---
    _add_signature_page(doc, metadata)

    # --- Serialisation ---
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
