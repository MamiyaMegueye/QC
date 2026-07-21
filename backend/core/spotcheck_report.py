"""
spotcheck_report.py — Generateur de rapport analytique Spotcheck enrichi IA.

Structure inspiree des rapports SISTA (bureau d'etudes Mauritanie) :
  - Page de titre
  - Resume executif (redige par IA)
  - I. Contexte et justification
  - II. Methodologie (Wilaya/Moughataa/Commune)
  - III. Deroulement (formation, phase collecte, problemes)
  - IV. Resultats en 5 sections analytiques (redigees par IA)
      4.1 Caracteristiques sociodemographiques
      4.2 Paiement des beneficiaires
      4.3 Utilisation du cash
      4.4 Inscription et connaissance du programme
      4.5 Reclamations des beneficiaires
  - V. Conclusion et recommandations (rediges par IA)

L'IA utilise Claude API pour rediger les interpretations qualitatives.
Si l'IA echoue (pas de cle, timeout), le rapport se degrade gracieusement.
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from datetime import datetime
import os
import json
import tempfile

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .spotcheck import (
    analyze_spotcheck, SECTION_LABELS, SECTION_DUREE_BORNES,
    parse_ts_unix, parse_duree_string
)

NAVY = '1E3A5F'
GOLD = 'C9A032'
RED = 'C41E3A'
ORANGE = 'F39C12'
GREEN = '27AE60'


# ============================================================
#  Helpers docx
# ============================================================

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading_styled(doc, text, level, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


# ============================================================
#  IA — Claude API pour redaction des interpretations
# ============================================================

def call_claude_for_report(section_id, context_data, api_key=None):
    """Appelle Claude pour rediger un texte d'interpretation.

    Args:
        section_id : 'resume_executif' | 'section_4_1' ... 'section_4_5' | 'conclusion'
        context_data : dict avec les stats a interpreter
        api_key : cle Anthropic (par defaut : ANTHROPIC_API_KEY env)

    Retourne le texte redige, ou None si echec.
    Le rapport reste generable meme si l'IA echoue.
    """
    api_key = api_key or os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("API1_KEY")
    if not api_key:
        return None

    try:
        from anthropic import Anthropic
    except ImportError:
        return None

    prompts = {
        "resume_executif": _prompt_resume_executif,
        "section_4_1": _prompt_section_1_sociodemo,
        "section_4_2": _prompt_section_2_paiement,
        "section_4_3": _prompt_section_3_cash,
        "section_4_4": _prompt_section_4_inscription,
        "section_4_5": _prompt_section_5_reclamations,
        "conclusion": _prompt_conclusion,
    }

    if section_id not in prompts:
        return None

    prompt_text = prompts[section_id](context_data)

    try:
        client = Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=1500,
            system=(
                "Tu es un consultant analyste travaillant pour le bureau d'etudes SISTA "
                "en Mauritanie, specialise dans l'evaluation de programmes de protection "
                "sociale. Tu rediges un rapport d'analyse Spotcheck du programme Tekavoul "
                "(transferts sociaux monetaires, PASyFiS II, Banque Mondiale / Taazour). "
                "\n\n"
                "Ton style : sobre, factuel, oriente decideurs. Tu utilises des chiffres "
                "concrets et tu formules des interpretations claires (pas juste des "
                "constats). Tu identifies les points forts et les problemes, avec des "
                "comparaisons entre editions ou entre equipes quand pertinent. "
                "\n\n"
                "Reponds en francais, en 2-4 paragraphes courts, sans titres ni listes. "
                "Ecris directement le texte comme s'il allait dans le rapport final."
            ),
            messages=[{"role": "user", "content": prompt_text}]
        )
        return message.content[0].text.strip()
    except Exception as e:
        print(f"[Claude error] {section_id}: {e}")
        return None


# ============================================================
#  Prompts specifiques par section
# ============================================================

def _prompt_resume_executif(ctx):
    return f"""Redige le RESUME EXECUTIF de notre rapport Spotcheck.

DONNEES DE L'ENQUETE :
- {ctx['n_total']} menages interviewes
- Periode : {ctx.get('periode', 'periode indiquee')}
- Zone : {ctx.get('zone', 'Nouakchott')}
- {ctx['n_enqueteurs']} enqueteurs sous supervision de {ctx['n_superviseurs']} superviseurs
- Duree mediane des interviews : {ctx['duree_med_min']} min (moyenne : {ctx['duree_moy_min']} min)

REPARTITION GEOGRAPHIQUE :
{ctx['geo_txt']}

ALERTES QUALITE DETECTEES :
- {ctx['n_interviews_baclees']} interviews sous les seuils min de sections ({ctx['pct_baclees']}%)
- {ctx['n_interviews_deviees']} interviews au-dessus des seuils max ({ctx['pct_deviees']}%)
- Coherence chronologique : {ctx['n_incoherences']} anomalies detectees

Redige un resume executif de 3-4 paragraphes couvrant :
1. Ce qui a ete fait (echantillon, methodologie, couverture)
2. Les principaux constats positifs
3. Les points d'attention principaux (par ordre de gravite)
4. Une conclusion d'une phrase sur la qualite globale de la collecte.
"""


def _prompt_section_1_sociodemo(ctx):
    return f"""Redige la sous-section 4.1 CARACTERISTIQUES SOCIODEMOGRAPHIQUES.

DONNEES DE LA SECTION 1 (Identification & sociodemo) :
{ctx['stats_txt']}

DUREE DE CETTE SECTION :
- Duree mediane : {ctx.get('section_med_sec', '?')}s (attendu : {ctx.get('section_cible_sec', '?')}s, seuils {ctx.get('borne_min', '?')}-{ctx.get('borne_max', '?')}s)
- {ctx.get('n_court', 0)} interviews trop courtes ({ctx.get('pct_court', 0)}%)
- {ctx.get('n_long', 0)} interviews trop longues ({ctx.get('pct_long', 0)}%)

Redige 3-4 paragraphes analytiques sur les caracteristiques sociodemographiques des
beneficiaires. Interprete les chiffres (proportion de femmes, langue majoritaire,
possession de telephone, etc.), pas juste les enumerer. Mentionne la qualite de la
collecte sur cette section si pertinent.
"""


def _prompt_section_2_paiement(ctx):
    return f"""Redige la sous-section 4.2 PAIEMENT DES BENEFICIAIRES.

DONNEES DE LA SECTION 2 (Paiement & transferts) :
{ctx['stats_txt']}

DUREE DE CETTE SECTION (LA PLUS CRITIQUE - c'est ici qu'on verifie le cash) :
- Duree mediane : {ctx.get('section_med_sec', '?')}s (attendu : {ctx.get('section_cible_sec', '?')}s, seuils {ctx.get('borne_min', '?')}-{ctx.get('borne_max', '?')}s)
- {ctx.get('n_court', 0)} interviews trop courtes ({ctx.get('pct_court', 0)}%) — signal potentiel de baclage sur la verification du paiement
- {ctx.get('n_long', 0)} interviews trop longues ({ctx.get('pct_long', 0)}%)

Redige 3-4 paragraphes couvrant :
- Taux de reception du cash (99,2% attendu au global)
- Canal d'information (SMS, voisins, etc.)
- Assistance au retrait
- Temps d'attente
- Satisfaction
- **Point d'alerte cle** : les frais de retrait declares malgre la gratuite prevue
- Distance percue

Interprete les chiffres. Cette section est le coeur du controle : si elle est bâclee,
c'est ici qu'il faut faire un back check terrain.
"""


def _prompt_section_3_cash(ctx):
    return f"""Redige la sous-section 4.3 UTILISATION DU CASH.

DONNEES DE LA SECTION 3 :
{ctx['stats_txt']}

DUREE :
- {ctx.get('section_med_sec', '?')}s de mediane (seuils {ctx.get('borne_min', '?')}-{ctx.get('borne_max', '?')}s)
- {ctx.get('n_court', 0)} courtes / {ctx.get('n_long', 0)} longues

Redige 2-3 paragraphes sur l'utilisation du cash recu et la perception des prix
(inflation eventuelle sur riz, huile, sucre, legumineuses).
"""


def _prompt_section_4_inscription(ctx):
    return f"""Redige la sous-section 4.4 INSCRIPTION ET CONNAISSANCE DU PROGRAMME.

DONNEES DE LA SECTION 4 :
{ctx['stats_txt']}

DUREE :
- {ctx.get('section_med_sec', '?')}s de mediane (seuils {ctx.get('borne_min', '?')}-{ctx.get('borne_max', '?')}s)
- {ctx.get('n_court', 0)} courtes / {ctx.get('n_long', 0)} longues

Redige 2-3 paragraphes sur :
- Sensibilisation lors de l'inscription (via animateurs sociaux)
- Connaissance des objectifs du programme
- Comprehension du lien cash-promotion sociale
- Explications sur Sedad

Si tu as des donnees d'une edition precedente (2025), fais la comparaison.
"""


def _prompt_section_5_reclamations(ctx):
    return f"""Redige la sous-section 4.5 RECLAMATIONS DES BENEFICIAIRES.

DONNEES DE LA SECTION 5 :
{ctx['stats_txt']}

DUREE :
- {ctx.get('section_med_sec', '?')}s de mediane (seuils {ctx.get('borne_min', '?')}-{ctx.get('borne_max', '?')}s)
- {ctx.get('n_court', 0)} courtes / {ctx.get('n_long', 0)} longues

Redige 2-3 paragraphes sur :
- Connaissance du numero vert Tekavoul
- Recours au mecanisme de plaintes
- Types de reclamations formulees
- Satisfaction du retour
"""


def _prompt_conclusion(ctx):
    return f"""Redige la CONCLUSION ET RECOMMANDATIONS du rapport.

SYNTHESE DES CONSTATS :
- Duree mediane : {ctx['duree_med_min']} min sur {ctx['n_total']} menages
- {ctx['n_interviews_baclees']} interviews baclees ({ctx['pct_baclees']}%)
- {ctx['n_interviews_deviees']} interviews au-dessus des seuils max
- Coherence : {ctx['n_incoherences']} problemes chronologiques
- {ctx['n_enqueteurs']} enqueteurs / {ctx['n_superviseurs']} superviseurs

ENQUETEURS LES PLUS SUSPECTS (pct de sections courtes) :
{ctx.get('enq_suspects_txt', '(pas d enqueteur particulierement suspect)')}

Redige 2 sous-sections :

**CONCLUSION** (2 paragraphes) : ce qui va bien, ce qui inquiete.

**RECOMMANDATIONS** (bullet points) : 5-8 actions concretes prioritaires, du plus
urgent au moins urgent. Focus sur :
- Actions immediates de back check
- Ameliorations du protocole de collecte
- Formation complementaire des enqueteurs identifies
- Configuration technique (XLSForm, GPS)
- Communication vers les beneficiaires (frais Sedad, numero vert)
"""


# ============================================================
#  Preparation des contextes pour l'IA
# ============================================================

def build_context_resume(df, analysis, metadata):
    n_total = len(df)
    matched = analysis["detection"]["matched_core"]
    n_enq = df[matched["enqueteur"]].nunique() if "enqueteur" in matched else 0
    n_sup = df[matched.get("superviseur")].nunique() if "superviseur" in matched else 0

    tot = analysis["durations"]["total"] or {}
    sections = analysis["durations"]["sections"]

    n_baclees = 0
    n_deviees = 0
    for s in sections:
        n_baclees += s.get("n_court", 0)
        n_deviees += s.get("n_long", 0)

    coh = analysis["coherence"]["counts"] if analysis["coherence"]["available"] else {}
    n_incoh = sum(coh.values())

    geo_txt = ""
    if analysis["geography"]["available"]:
        for w in analysis["geography"]["wilayas"]:
            geo_txt += f"- {w['name']} : {w['n_menages']} menages ({w['pct']}%)\n"

    return {
        "n_total": n_total,
        "n_enqueteurs": n_enq,
        "n_superviseurs": n_sup,
        "duree_med_min": tot.get("med_min", 0),
        "duree_moy_min": tot.get("moy_min", 0),
        "geo_txt": geo_txt,
        "n_interviews_baclees": n_baclees,
        "pct_baclees": round(n_baclees / max(1, n_total * len(sections)) * 100, 1),
        "n_interviews_deviees": n_deviees,
        "pct_deviees": round(n_deviees / max(1, n_total * len(sections)) * 100, 1),
        "n_incoherences": n_incoh,
        "periode": metadata.get("periode", ""),
        "zone": metadata.get("zone", "Nouakchott"),
    }


def build_context_section(df, section_num, analysis):
    """Prepare le contexte pour l'IA pour une section 4.X donnee."""
    # Trouver les colonnes de la section
    section_prefixes = {
        1: ["I1", "I2", "I3", "I4", "I5", "I6", "S1"],
        2: ["II1", "II2", "II3", "II9", "II10", "II11", "II12", "II13",
            "II14", "II15", "II16", "II17", "II18", "II19", "II20", "II21", "CN1"],
        3: ["III1", "III2", "III3"],
        4: ["IV1", "IV2", "IV7", "IV8", "IV9"],
        5: ["V2", "V3", "V4", "V5", "V6", "V7", "V8"],
    }
    prefixes = section_prefixes.get(section_num, [])
    cols_section = []
    for col in df.columns:
        col_norm = str(col).upper().replace(".", "").replace("_", "").replace(" ", "")
        for pref in prefixes:
            if col_norm.startswith(pref):
                cols_section.append(col)
                break

    stats_txt = ""
    for col in cols_section[:12]:  # limite a 12 variables
        s = df[col].dropna()
        n_valid = len(s)
        if n_valid == 0:
            continue
        try:
            n_uniques = s.nunique()
        except TypeError:
            continue
        if n_uniques <= 8 and n_uniques > 0:
            counts = s.astype(str).value_counts()
            top = ", ".join([
                f"{v}={counts[v]} ({counts[v]/n_valid*100:.1f}%)"
                for v in counts.head(4).index
            ])
            stats_txt += f"- {col} (n={n_valid}): {top}\n"
        elif pd.api.types.is_numeric_dtype(s):
            try:
                stats_txt += f"- {col} (n={n_valid}): moy={float(s.mean()):.1f}, med={float(s.median()):.1f}\n"
            except Exception:
                pass

    # Duree de la section
    section_data = next(
        (s for s in analysis["durations"]["sections"] if s["num"] == section_num),
        {}
    )

    return {
        "stats_txt": stats_txt or "(pas de donnees exploitables sur cette section)",
        "section_med_sec": section_data.get("med_sec", 0),
        "section_cible_sec": section_data.get("cible", 0),
        "borne_min": section_data.get("borne_min", 0),
        "borne_max": section_data.get("borne_max", 0),
        "n_court": section_data.get("n_court", 0),
        "pct_court": section_data.get("pct_court", 0),
        "n_long": section_data.get("n_long", 0),
        "pct_long": section_data.get("pct_long", 0),
    }


def build_context_conclusion(df, analysis, metadata, base_ctx):
    """Contexte pour la conclusion : ajoute les enqueteurs suspects."""
    ctx = dict(base_ctx)
    enq_data = analysis["by_enqueteur"]
    # Top 5 enqueteurs avec le plus de sections courtes
    scored = []
    for e in enq_data:
        total_court = sum(s.get("pct_court", 0) for s in e["sections"])
        scored.append((e["enqueteur"], e["n_interviews"], total_court, e["sections"]))
    scored.sort(key=lambda x: -x[2])

    enq_txt = ""
    for enq, n, score, sections in scored[:5]:
        details = []
        for s in sections:
            if s.get("pct_court", 0) > 15:
                details.append(f"S{s['num']}:{s['pct_court']}%")
        if details:
            enq_txt += f"- ENQ {enq} (n={n}) : {', '.join(details)}\n"

    ctx["enq_suspects_txt"] = enq_txt if enq_txt else "(pas d'enqueteur particulierement suspect)"
    return ctx


# ============================================================
#  Graphiques
# ============================================================

def make_chart_dir():
    return tempfile.mkdtemp(prefix="spotcheck_charts_")


def chart_wilaya_pie(geo_data, chart_dir):
    if not geo_data.get("available"):
        return None
    wilayas = geo_data["wilayas"]
    if not wilayas:
        return None
    labels = [w["name"] for w in wilayas]
    sizes = [w["n_menages"] for w in wilayas]
    colors = [f'#{NAVY}', f'#{GOLD}', '#4A6FA5', '#8B7B3E', '#2C3E50']

    fig, ax = plt.subplots(figsize=(6.5, 5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct='%1.1f%%',
        colors=colors[:len(sizes)], startangle=90,
        wedgeprops=dict(edgecolor='white', linewidth=2),
        textprops=dict(fontsize=11),
    )
    for at in autotexts:
        at.set_color('white')
        at.set_fontweight('bold')
    ax.set_title("Repartition de l'echantillon par Wilaya",
                 fontsize=13, fontweight='bold', pad=15)
    plt.tight_layout()
    path = os.path.join(chart_dir, "wilaya_pie.png")
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def chart_section_durations(sections_data, chart_dir):
    if not sections_data.get("available"):
        return None
    sections = sections_data["sections"]
    if not sections:
        return None

    labels = [f"S{s['num']}\n{s['label'][:25]}" for s in sections]
    med_min = [s['med_sec'] / 60 for s in sections]
    borne_min = [s['borne_min'] / 60 for s in sections]
    borne_max = [s['borne_max'] / 60 for s in sections]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    x = np.arange(len(labels))
    bars = ax.bar(x, med_min, color=f'#{NAVY}', edgecolor='white',
                  alpha=0.85, label='Mediane observee')

    # Zone verte des bornes acceptees
    for i, (bmin, bmax) in enumerate(zip(borne_min, borne_max)):
        ax.hlines(bmin, i - 0.4, i + 0.4, colors='green', linestyles='dashed', linewidth=1)
        ax.hlines(bmax, i - 0.4, i + 0.4, colors='red', linestyles='dashed', linewidth=1)

    for i, m in enumerate(med_min):
        ax.text(i, m + 0.1, f'{m:.1f} min', ha='center', va='bottom',
                fontsize=9, fontweight='bold')

    ax.set_ylabel('Duree (minutes)', fontsize=11)
    ax.set_title('Duree mediane par section (avec bornes min/max)',
                 fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.grid(True, alpha=0.3, axis='y')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(chart_dir, "section_durations.png")
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def chart_sections_hors_bornes(sections_data, chart_dir):
    """Bar chart empile : par section, % court + % long."""
    if not sections_data.get("available"):
        return None
    sections = sections_data["sections"]
    if not sections:
        return None

    labels = [f"S{s['num']}" for s in sections]
    pct_court = [s.get('pct_court', 0) for s in sections]
    pct_long = [s.get('pct_long', 0) for s in sections]

    fig, ax = plt.subplots(figsize=(8, 4))
    x = np.arange(len(labels))
    width = 0.35
    ax.bar(x - width/2, pct_court, width, label='% trop court', color=f'#{ORANGE}', edgecolor='white')
    ax.bar(x + width/2, pct_long, width, label='% trop long', color=f'#{RED}', edgecolor='white')

    for i, (c, l) in enumerate(zip(pct_court, pct_long)):
        if c > 0:
            ax.text(i - width/2, c + 0.3, f'{c}%', ha='center', va='bottom', fontsize=9)
        if l > 0:
            ax.text(i + width/2, l + 0.3, f'{l}%', ha='center', va='bottom', fontsize=9)

    ax.set_ylabel('% des interviews', fontsize=11)
    ax.set_title('Sections hors bornes par type', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels([f"S{s['num']} - {s['label'][:20]}" for s in sections], fontsize=9)
    ax.legend()
    ax.grid(True, alpha=0.3, axis='y')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(chart_dir, "sections_hors_bornes.png")
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


# ============================================================
#  Generation du rapport complet
# ============================================================

def generate_spotcheck_report(df, output_path, metadata=None, api_key=None):
    """Genere le rapport analytique Spotcheck avec interpretations IA.

    Args:
        df : DataFrame des donnees
        output_path : chemin du .docx a generer
        metadata : dict optionnel (programme, projet, bureau, zone, periode)
        api_key : cle Anthropic (optionnel, sinon prise dans l'env)

    Returns:
        Le path du .docx genere.
    """
    metadata = metadata or {}
    analysis = analyze_spotcheck(df)
    n_total = len(df)

    if not analysis["available"]:
        raise ValueError(
            "Format non reconnu comme Spotcheck. "
            "Verifiez que le fichier contient CODE, WILAYA, SUPERVISEUR, ENQUETEUR "
            "et au moins DS1/FS1/DUREES1."
        )

    chart_dir = make_chart_dir()

    # Preparer les graphiques
    wilaya_chart = chart_wilaya_pie(analysis["geography"], chart_dir)
    section_chart = chart_section_durations(analysis["durations"], chart_dir)
    bornes_chart = chart_sections_hors_bornes(analysis["durations"], chart_dir)

    # Preparer les contextes et appeler l'IA
    ctx_resume = build_context_resume(df, analysis, metadata)
    ai_resume = call_claude_for_report("resume_executif", ctx_resume, api_key)

    ai_sections = {}
    for section_num in range(1, min(6, analysis["detection"]["n_sections"] + 1)):
        ctx = build_context_section(df, section_num, analysis)
        ai_text = call_claude_for_report(f"section_4_{section_num}", ctx, api_key)
        ai_sections[section_num] = ai_text

    ctx_conclusion = build_context_conclusion(df, analysis, metadata, ctx_resume)
    ai_conclusion = call_claude_for_report("conclusion", ctx_conclusion, api_key)

    # ============================================================
    #  Construction du document
    # ============================================================
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- PAGE DE TITRE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(metadata.get("projet", "Projet d'Appui au Systeme de Filets Sociaux II (PASyFiS II)"))
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_paragraph()
    title_table = doc.add_table(rows=1, cols=1)
    title_table.style = 'Light Grid Accent 1'
    tc = title_table.rows[0].cells[0]
    tc_p = tc.paragraphs[0]
    tc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tc_p.add_run('RAPPORT D\'ANALYSE\n')
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r = tc_p.add_run(f"\nEnquete de Verification Ponctuelle (Spot Check)\n")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r = tc_p.add_run(f"\nProgramme {metadata.get('programme', 'Tekavoul')} Regulier\n")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{metadata.get('zone', 'Nouakchott')} — {metadata.get('periode', datetime.now().strftime('%B %Y'))}")
    r.font.size = Pt(14)
    r.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n" + metadata.get("bureau", "SISTA Consult"))
    r.font.size = Pt(11)
    r.italic = True

    doc.add_page_break()

    # --- RESUME EXECUTIF ---
    add_heading_styled(doc, "Resume executif", 1)
    if ai_resume:
        for para in ai_resume.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        # Fallback deterministe
        p = doc.add_paragraph()
        p.add_run(
            f"La presente verification ponctuelle a touche "
        )
        r = p.add_run(f"{n_total} menages")
        r.bold = True
        p.add_run(f" sur la periode {metadata.get('periode', 'indiquee')}. ")
        if analysis["geography"]["available"]:
            wilayas = analysis["geography"]["wilayas"]
            p.add_run(f"{len(wilayas)} Wilaya(s) ont ete couvertes.")

    doc.add_page_break()

    # --- I. CONTEXTE ---
    add_heading_styled(doc, "I. Contexte et justification", 1)
    p = doc.add_paragraph()
    p.add_run(
        "La Republique Islamique de Mauritanie a entrepris des efforts considerables "
        "pour renforcer son systeme national de protection sociale. Dans ce cadre, "
        "le Projet d'Appui au Systeme de Filets Sociaux II (PASyFiS II), mis en oeuvre "
        "par la Delegation Generale a la Solidarite Nationale et a la Lutte contre "
        "l'Exclusion (Taazour) avec l'appui de la Banque mondiale, s'appuie sur le "
        f"programme {metadata.get('programme', 'Tekavoul')} pour distribuer des "
        "transferts monetaires reguliers aux menages beneficiaires."
    )
    p = doc.add_paragraph()
    p.add_run(
        "Les spot checks constituent un outil essentiel de suivi et d'assurance qualite. "
        "Ils permettent de verifier si les activites menees respectent les procedures "
        "prevues dans les manuels operationnels, d'apprecier le niveau de comprehension "
        "des processus par les beneficiaires, et de mesurer leur degre de satisfaction."
    )

    # --- II. METHODOLOGIE ---
    add_heading_styled(doc, "II. Methodologie", 1)
    add_heading_styled(doc, "Taille et allocation de l'echantillon", 2)

    if analysis["geography"]["available"]:
        wilayas = analysis["geography"]["wilayas"]

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, h in enumerate(['Wilaya', 'Nb Moughataas', 'Nb menages', 'Pourcentage']):
            hdr[i].text = h
            for pp in hdr[i].paragraphs:
                for rn in pp.runs:
                    rn.font.bold = True

        for w in wilayas:
            r = table.add_row().cells
            r[0].text = w["name"]
            r[1].text = str(len(w["moughataas"]))
            r[2].text = str(w["n_menages"])
            r[3].text = f"{w['pct']} %"

        r = table.add_row().cells
        r[0].text = 'TOTAL'
        r[1].text = str(sum(len(w["moughataas"]) for w in wilayas))
        r[2].text = str(n_total)
        r[3].text = '100 %'
        for c in r:
            for pp in c.paragraphs:
                for rn in pp.runs:
                    rn.font.bold = True

        if wilaya_chart:
            doc.add_paragraph()
            doc.add_picture(wilaya_chart, width=Inches(4.5))

    # --- III. DEROULEMENT ---
    add_heading_styled(doc, "III. Deroulement de l'enquete", 1)

    matched = analysis["detection"]["matched_core"]
    n_sup = df[matched.get("superviseur")].nunique() if "superviseur" in matched else 0
    n_enq = df[matched["enqueteur"]].nunique() if "enqueteur" in matched else 0

    p = doc.add_paragraph()
    p.add_run("La collecte a mobilise ")
    r = p.add_run(f"{n_enq} enqueteurs")
    r.bold = True
    p.add_run(" repartis sous la supervision de ")
    r = p.add_run(f"{n_sup} superviseurs")
    r.bold = True
    p.add_run(
        ". Un systeme de suivi journalier a ete mis en place avec un controle "
        "qualite en temps reel analysant les bases recues chaque jour : codes saisis, "
        "valeurs incoherentes, temps d'interview anormaux et localisations GPS suspectes."
    )

    add_heading_styled(doc, "Duree des interviews par section", 2)
    if section_chart:
        doc.add_picture(section_chart, width=Inches(6.3))

    p = doc.add_paragraph()
    p.add_run(
        "Chaque section du questionnaire a une duree attendue definie par le protocole "
        "de collecte. Une section trop courte peut signaler un baclage, tandis qu'une "
        "section anormalement longue peut indiquer une deviation du protocole ou un "
        "remplissage manuel apres coup."
    )

    if bornes_chart:
        doc.add_paragraph()
        doc.add_picture(bornes_chart, width=Inches(6.3))

    # Tableau detaille par section
    sections = analysis["durations"]["sections"]
    if sections:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, h in enumerate(['Section', 'Libelle', 'Med. (s)', 'Bornes (s)',
                                '<min', '>max', 'Verdict']):
            hdr[i].text = h
            for pp in hdr[i].paragraphs:
                for rn in pp.runs:
                    rn.font.bold = True

        for s in sections:
            r = table.add_row().cells
            r[0].text = f"S{s['num']}"
            r[1].text = s["label"][:35]
            r[2].text = f"{s['med_sec']:.0f}"
            r[3].text = f"{s['borne_min']}-{s['borne_max']}"
            r[4].text = f"{s['n_court']} ({s['pct_court']}%)"
            r[5].text = f"{s['n_long']} ({s['pct_long']}%)"

            total_hors = s['pct_court'] + s['pct_long']
            if total_hors > 30:
                r[6].text = "ALERTE"
                set_cell_bg(r[6], RED)
            elif total_hors > 15:
                r[6].text = "Attention"
                set_cell_bg(r[6], ORANGE)
            else:
                r[6].text = "OK"
                set_cell_bg(r[6], GREEN)
            for pp in r[6].paragraphs:
                for rn in pp.runs:
                    rn.font.color.rgb = RGBColor.from_string('FFFFFF')
                    rn.font.bold = True

    doc.add_page_break()

    # --- IV. RESULTATS ---
    add_heading_styled(doc, "IV. Resultats de l'enquete", 1)
    p = doc.add_paragraph()
    p.add_run(
        f"L'echantillon de l'enquete est compose de {n_total} repondants. "
        "L'analyse s'articule autour de cinq sections respectant la structure logique "
        "du questionnaire : (i) caracteristiques sociodemographiques, (ii) mode de "
        "paiement et satisfaction du processus de retrait, (iii) utilisation du cash "
        "recu, (iv) inscription et connaissance du programme, (v) mode de reclamation."
    )

    section_titles = {
        1: "4.1 Caracteristiques sociodemographiques",
        2: "4.2 Paiement des beneficiaires",
        3: "4.3 Utilisation du cash",
        4: "4.4 Inscription et connaissance du programme",
        5: "4.5 Reclamations des beneficiaires",
    }

    for section_num in sorted(section_titles.keys()):
        if section_num > analysis["detection"]["n_sections"]:
            break
        add_heading_styled(doc, section_titles[section_num], 2)
        ai_text = ai_sections.get(section_num)
        if ai_text:
            for para in ai_text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            # Fallback deterministe basique
            p = doc.add_paragraph()
            p.add_run(
                f"[Interpretation IA non disponible pour cette section. "
                f"Consultez les tableaux statistiques detailles en annexe.]"
            ).italic = True

    doc.add_page_break()

    # --- V. CONCLUSION ---
    add_heading_styled(doc, "V. Conclusion et recommandations", 1)
    if ai_conclusion:
        for para in ai_conclusion.split("\n\n"):
            if para.strip():
                # Detecter si c'est un bullet
                if para.strip().startswith(("- ", "• ", "* ")):
                    for line in para.strip().split("\n"):
                        clean = line.lstrip("-•* ").strip()
                        if clean:
                            doc.add_paragraph(clean, style='List Bullet')
                else:
                    doc.add_paragraph(para.strip())
    else:
        p = doc.add_paragraph()
        p.add_run(
            f"L'enquete Spotcheck realisee sur {n_total} menages a permis d'analyser "
            "en profondeur la mise en oeuvre du programme et de detecter des anomalies "
            "qualitatives qui appellent des mesures correctives."
        )

    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'Rapport genere le {datetime.now().strftime("%d/%m/%Y a %H:%M")}').italic = True
    if ai_resume:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run('Interpretations qualitatives assistees par IA (Claude)').italic = True

    doc.save(output_path)

    try:
        import shutil
        shutil.rmtree(chart_dir, ignore_errors=True)
    except Exception:
        pass

    return output_path